import os
from tkinter import *
import re
from tkinter.filedialog import *
from tkinter import messagebox
import tkinter.scrolledtext as tkscrolled

import time
from datetime import datetime
import csv

from docx import Document
from docx.shared import Inches
from docx.blkcntnr import BlockItemContainer
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.section import Section, Sections
from docx.shared import ElementProxy, Emu
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_ROW_HEIGHT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Border, Side, Alignment, Protection, Font, NamedStyle
from openpyxl.styles.colors import RED

class Felulet():
    def __init__(self, master):
        self.count_auto=0
        self.ellatasok=[]
        self.ugyfellista={}
        self.ugyfellista_nappali={}
        self.ugyfellista_ejjeli ={}

        self.path=""
        frame1 = Frame(master)
        self.menu = Menu(master)
        self.iranyitoszamok={}
        self.iranyitoszamok_betoltese()
        self.korzetszamok={}
        self.korzetszamok_betoltese()
        master.config(menu=self.menu)
        self.person_number=IntVar()
        self.person_number.set(1)

        self.person_number_new=1

        self.error=False
        self.changed=False
        self.master = master
        self.filemenu = Menu()
        self.filemenu.add_command(label="Új adatbázis létrehozása", command=self.new)
        self.filemenu.add_command(label="Adatbázis megnyitása", command=self.open)
        self.filemenu.add_command(label="Adatbázis mentése", command=self.save)
        self.filemenu.add_command(label="Adatbázis mentése másként", command=self.save_as)
        self.filemenu.add_command(label="Adatbázis exportálása Excel formátumban", command=self.export1)
        self.filemenu.add_command(label="Ügyfél adatainak exportálása adatlapként", command=self.export2)
        self.filemenu.add_command(label="Csoportos adatlap kimentés", command=self.export3)

        self.filemenu.add_separator()
        self.filemenu.add_command(label="Kilépés", command=self.close)
        self.menu.add_cascade(label="Fájl", menu=self.filemenu)
        self.menu.add_cascade(label="Ügyfél adatlapjának mentése", command=self.save_person)
        self.menu.add_cascade(label="Adatlap ürítése", command=self.zero_record)
        self.menu.add_cascade(label="Keresés az ügyfelek között", command=self.search)
        self.menu.add_cascade(label="Névjegy", command=self.nevjegy)

        self.scroller = Scrollbar(master)
        self.list_of_people=Listbox(master, yscrollcommand=self.scroller.set, width=30)
        self.scroller.config(command=self.list_of_people.yview)
        self.active_name = self.list_of_people.get(ACTIVE)
        self.now=self.list_of_people.curselection()

        eltol=100
        self.list_of_people.pack(side=LEFT, fill=Y, anchor=W)
        self.scroller.pack(side=LEFT, fill=BOTH, anchor=W)
        self.text1=Label(master, text="Sorszám:")
        self.text1.place(x=215, y=55)
        self.number_entry=Entry(master, textvariable=self.person_number, width=3, state=DISABLED, justify=RIGHT)
        self.number_entry.place(x=215, y=75)
        self.text2=Label(master, text="Ügyfél neve:")
        self.text2.place(x=310+eltol, y=5)
        self.name_var=StringVar()
        self.name_entry=Entry(master, textvariable=self.name_var, width=20)
        self.name_entry.place(x=385+eltol, y=5)
        self.text3=Label(master, text="TAJ szám:")
        self.text3.place(x=520+eltol, y=5)
        self.TAJ1_V=StringVar()
        self.TAJ2_V=StringVar()
        self.TAJ3_V=StringVar()
        self.TAJ4_V=StringVar()

        self.TAJ1_Entry=Entry(master, textvariable=self.TAJ1_V, width=3, justify=RIGHT)
        self.TAJ2_Entry=Entry(master, textvariable=self.TAJ2_V, width=3, justify=RIGHT)
        self.TAJ3_Entry=Entry(master, textvariable=self.TAJ3_V, width=2, justify=RIGHT)
        self.TAJ4_Entry=Entry(master, textvariable=self.TAJ4_V, width=1, justify=RIGHT, state=DISABLED)

        self.TAJ1_Entry.place(x=580+eltol, y=5)
        self.TAJ_Line1=Label(master, text="-")
        self.TAJ_Line1.place(x=605+eltol, y=5)
        self.TAJ2_Entry.place(x=620+eltol, y=5)
        self.TAJ_Line2=Label(master, text="-")
        self.TAJ_Line2.place(x=645+eltol, y=5)
        self.TAJ3_Entry.place(x=660+eltol, y=5)
        self.TAJ4_Entry.place(x=675+eltol, y=5)

        self.text4 = Label(master, text="Születési név:")
        self.text4.place(x=305+eltol, y=30)
        self.bname_var = StringVar()
        self.bname_entry = Entry(master, textvariable=self.bname_var, width=20)
        self.bname_entry.place(x=385+eltol, y=30)
        self.same_V=BooleanVar()
        self.same_V.set(1)
        self.nincsTAJ=BooleanVar()
        self.nincsTAJnincs=Checkbutton(master, variable=self.nincsTAJ, text="Nincs TAJ száma")
        self.nincsTAJnincs.place(x=518+eltol, y=25)
        self.same=Checkbutton(master, variable=self.same_V, text="A viselt névvel megegyező")
        self.same.place(x=518+eltol, y=45)
        self.mothername=StringVar()
        self.text33=Label(master, text="Édesanyja neve:")
        self.mothername_entry=Entry(master, textvariable=self.mothername, width=20)
        self.text33.place(x=285+eltol, y=52)
        self.mothername_entry.place(x=385+eltol, y=52)
        self.text5=Label(master, text="Állampolgárság:")
        self.text5.place(x=290+eltol, y=75)
        self.nation_var=IntVar()
        self.nation_var.set(1)
        self.nation_button=Radiobutton(master, variable=self.nation_var, value=1, text="Magyar")
        self.nation_button.place(x=380+eltol, y=75)
        self.nation_button2 = Radiobutton(master, variable=self.nation_var, value=2, text="Egyéb:")
        self.nation_button2.place(x=450+eltol, y=75)
        self.nationality_var=StringVar()
        self.nationality_entry=Entry(master, textvariable=self.nationality_var, width=15)
        self.nationality_entry.place(x=522+eltol, y=75)
        self.text6=Label(master, text="Születési hely:")
        self.text6.place(x=302+eltol, y=135)
        self.bplace_V=StringVar()
        self.bplace = Entry(master, textvariable=self.bplace_V, width=20)
        self.bplace.place(x=385+eltol, y=135)
        self.text7=Label(master, text="Születési idő:")
        self.text7.place(x=518+eltol, y=135)
        self.bdate1_V=StringVar()
        self.bdate2_V=StringVar()
        self.bdate3_V=StringVar()
        self.bdate1_E=Entry(master, textvariable=self.bdate1_V, width=4, justify=RIGHT)
        self.bdate2_E=Entry(master, textvariable=self.bdate2_V, width=2, justify=RIGHT)
        self.bdate3_E=Entry(master, textvariable=self.bdate3_V, width=2, justify=RIGHT)
        self.bdate1_L=Label(master, text="-")
        self.bdate2_L=Label(master, text="-")
        self.bdate1_E.place(x=595+eltol, y=135)
        self.bdate1_L.place(x=625+eltol, y=135)
        self.bdate2_E.place(x=635+eltol, y=135)
        self.bdate2_L.place(x=650+eltol, y=135)
        self.bdate3_E.place(x=662+eltol, y=135)
        self.phone_number1=StringVar()
        self.phone_number2=StringVar()
        self.phone_number1V=Entry(master, width=2, justify=RIGHT, textvariable=self.phone_number1)
        self.phone_number2V=Entry(master, width=7, justify=RIGHT, textvariable=self.phone_number2)
        self.text7=Label(master,text="Telefonszám:")
        self.text8=Label(master,text="/")
        self.text7.place(x=516+eltol, y=105)
        self.phone_number1V.place(x=595+eltol, y=105)
        self.text8.place(x=615+eltol, y=105)
        self.phone_number2V.place(x=628+eltol, y=105)
        self.idegen=StringVar()
        self.idegens=["Bevándorolt", "Letelepedett", "Oltalmazott", "Menekült", "Menedékes"]
        self.idegens.sort()
        self.idegens=tuple(self.idegens)
        self.idegenT=Label(master, text="Idegenrendészeti státusz:")
        self.idegenT.place(x=240+eltol, y=104)
        self.idegenek = OptionMenu(master, self.idegen, *self.idegens)
        self.idegenek.place(x=383+eltol, y=100)
        self.cselekves=IntVar()
        self.cselekves.set(1)
        self.cselekvo1=Radiobutton(master, text="Cselekvőképes", variable=self.cselekves, value=1)
        self.cselekvo2=Radiobutton(master, text="Korlátozottan cselekvőképes", variable=self.cselekves, value=2)
        self.cselekvo3=Radiobutton(master, text="Cselekvőképtelen", variable=self.cselekves, value=3)
        self.text9=Label(master, text="Cselekvőképesség:")
        self.text9.place(x=250+eltol, y=160)
        self.cselekvo1.place(x=250+eltol, y=185)
        self.cselekvo2.place(x=360+eltol, y=185)
        self.cselekvo3.place(x=250+eltol, y=210)
        self.gondnok=BooleanVar()
        self.gondnokolt=Checkbutton(master, text="Gondnokolt", variable=self.gondnok)
        self.gondnokolt.place(x=440+eltol, y=160)
        self.text34=Label(master, text="Gondnok neve:")
        self.text34.place(x=540+eltol, y=162)
        self.text35=Label(master, text="Gondnok elérhetősége:")
        self.text35.place(x=405+eltol, y=212)

        self.gondnok_nev=StringVar()
        self.gondnok_entry=Entry(master, width=22, textvariable=self.gondnok_nev)
        self.gondnok_entry.place(x=542+eltol, y=187)
        self.gondnok_cim=StringVar()
        self.gondnok_cim_entry=Entry(master, width=22, textvariable=self.gondnok_cim)
        self.gondnok_cim_entry.place(x=542+eltol, y=212)

        self.text10=Label(master, text="Utolsó állandó lakhely:")
        self.text11=Label(master, text="Irányítószám:")
        self.text10.place(x=250+eltol, y=235)
        self.text11.place(x=387+eltol, y=235)
        self.IRSZ_1=StringVar()
        self.IRSZ_1E=Entry(master, width=4, textvariable=self.IRSZ_1, justify=RIGHT)
        self.IRSZ_1E.place(x=470+eltol, y=235)
        self.text12=Label(master, text="Város:")
        self.text12.place(x=425+eltol, y=255)
        self.city_1=StringVar()
        self.city_1E=Entry(master, width=30, textvariable=self.city_1)
        self.city_1E.place(x=470+eltol, y=255)
        self.ucca1=Label(master,text="Közterület neve:")
        self.ucca1.place(x=373+eltol, y=275)
        self.ucca1_V=StringVar()
        self.ucca1_E=Entry(master, width=30, textvariable=self.ucca1_V)
        self.ucca1_E.place(x=470+eltol, y=275)
        menuhoz=["", "utca", "út", "tér", "egyéb:"]
        self.text13=Label(master, text="Közterület típusa:")
        self.kozterulet_tipusa=StringVar()
        self.kozterulet_tipusa_menu=OptionMenu(master, self.kozterulet_tipusa, *menuhoz)
        self.kozterulet_tipusa_menu.config(width=6)
        self.text13.place(x=370+eltol, y=298)
        self.kozterulet_tipusa_menu.place(x=468+eltol, y=293)
        self.kozterulet_tipusa_egyeb=StringVar()
        self.kozterulet_tipusa_egyeb_Entry=Entry(master, width=16, textvariable=self.kozterulet_tipusa_egyeb, state=DISABLED)
        self.kozterulet_tipusa_egyeb_Entry.place(x=555+eltol, y=299)
        self.text22=Label(master, text="Házszám:")
        self.text22.place(x=370+eltol, y=320)
        self.address_V=StringVar()
        self.address_E=Entry(master, textvariable=self.address_V, width=31)
        self.address_E.place(x=470+eltol, y=320)

        self.text14 = Label(master, text="Tartózkodási hely:")
        self.text15 = Label(master, text="Irányítószám:")
        self.text14.place(x=250+eltol, y=350)
        self.text15.place(x=387+eltol, y=350)
        self.IRSZ_2 = StringVar()
        self.IRSZ_2E = Entry(master, width=4, textvariable=self.IRSZ_2, justify=RIGHT)
        self.IRSZ_2E.place(x=470+eltol, y=350)
        self.text16 = Label(master, text="Város:")
        self.text16.place(x=425+eltol, y=370)
        self.city_2 = StringVar()
        self.city_2E = Entry(master, width=30, textvariable=self.city_2)
        self.city_2E.place(x=470+eltol, y=370)
        self.ucca2 = Label(master, text="Közterület neve:")
        self.ucca2.place(x=373+eltol, y=390)
        self.ucca2_V = StringVar()
        self.ucca2_E = Entry(master, width=30, textvariable=self.ucca2_V)
        self.ucca2_E.place(x=470+eltol, y=390)
        self.text17 = Label(master, text="Közterület típusa:")
        self.kozterulet_tipusa2 = StringVar()
        self.kozterulet_tipusa2_menu = OptionMenu(master, self.kozterulet_tipusa2, *menuhoz)
        self.kozterulet_tipusa2_menu.config(width=6)
        self.text17.place(x=370+eltol, y=413)
        self.kozterulet_tipusa2_menu.place(x=468+eltol, y=408)
        self.kozterulet_tipusa2_egyeb = StringVar()
        self.kozterulet_tipusa2_egyeb_Entry = Entry(master, width=16, textvariable=self.kozterulet_tipusa2_egyeb,
                                                   state=DISABLED)
        self.kozterulet_tipusa2_egyeb_Entry.place(x=555+eltol, y=414)
        self.megegyezik1=BooleanVar()
        self.megegyezik1.set(1)
        self.megegyezik1_B=button=Checkbutton(master, text="Lakcímmel megegyezik", variable=self.megegyezik1)
        self.megegyezik1_B.place(x=505+eltol, y=350)
        self.text23=Label(master, text="Házszám:")
        self.text23.place(x=370+eltol, y=435)
        self.address2_V=StringVar()
        self.address2_E=Entry(master, textvariable=self.address2_V, width=31)
        self.address2_E.place(x=470+eltol, y=435)


        self.text18 = Label(master, text="Értesítési cím:")
        self.text19 = Label(master, text="Irányítószám:")
        self.text18.place(x=250+eltol, y=465)
        self.text19.place(x=387+eltol, y=465)
        self.IRSZ_3 = StringVar()
        self.IRSZ_3E = Entry(master, width=4, textvariable=self.IRSZ_3, justify=RIGHT)
        self.IRSZ_3E.place(x=470+eltol, y=465)
        self.text20 = Label(master, text="Város:")
        self.text20.place(x=425+eltol, y=485)
        self.city_3 = StringVar()
        self.city_3E = Entry(master, width=30, textvariable=self.city_3)
        self.city_3E.place(x=470+eltol, y=485)
        self.ucca3 = Label(master, text="Közterület neve:")
        self.ucca3.place(x=373+eltol, y=505)
        self.ucca3_V = StringVar()
        self.ucca3_E = Entry(master, width=30, textvariable=self.ucca3_V)
        self.ucca3_E.place(x=470+eltol, y=505)
        self.text21 = Label(master, text="Közterület típusa:")
        self.kozterulet_tipusa3 = StringVar()
        self.kozterulet_tipusa3_menu = OptionMenu(master, self.kozterulet_tipusa2, *menuhoz)
        self.kozterulet_tipusa3_menu.config(width=6)
        self.text21.place(x=370+eltol, y=528)
        self.kozterulet_tipusa3_menu.place(x=468+eltol, y=523)
        self.kozterulet_tipusa3_egyeb = StringVar()
        self.kozterulet_tipusa3_egyeb_Entry = Entry(master, width=16, textvariable=self.kozterulet_tipusa3_egyeb,
                                                   state=DISABLED)
        self.kozterulet_tipusa3_egyeb_Entry.place(x=555+eltol, y=529)
        self.megegyezik2=BooleanVar()
        self.megegyezik2.set(1)
        self.megegyezik2_B=button=Checkbutton(master, text="Tart. címmel megegyezik", variable=self.megegyezik2)
        self.megegyezik2_B.place(x=505+eltol, y=460)
        self.text24=Label(master, text="Házszám:")
        self.text24.place(x=370+eltol, y=550)
        self.address3_V=StringVar()
        self.address3_E=Entry(master, textvariable=self.address3_V, width=31)
        self.address3_E.place(x=470+eltol, y=550)
        self.text25=Label(master, text="Nagykorú hozzátartozó:")
        self.text25.place(x=335+eltol, y=575)
        self.hozzatart_V=StringVar()
        self.hozzatartozo_E=Entry(master, textvariable=self.hozzatart_V, width=31)
        self.hozzatartozo_E.place(x=470+eltol, y=575)
        self.text26=Label(master, text="Elérhetősége:")
        self.text26.place(x=390+eltol, y=600)
        self.eler_V=StringVar()
        self.eler_E=Entry(master, textvariable=self.eler_V, width=31)
        self.eler_E.place(x=470+eltol, y=600)
        self.text27=Label(master, text="Kérelem beterjesztése:")
        self.text27.place(x=344+eltol, y=620)
        self.kdate1_V = StringVar()
        self.kdate2_V = StringVar()
        self.kdate3_V = StringVar()
        self.kdate1_E = Entry(master, textvariable=self.kdate1_V, width=4, justify=RIGHT)
        self.kdate2_E = Entry(master, textvariable=self.kdate2_V, width=2, justify=RIGHT)
        self.kdate3_E = Entry(master, textvariable=self.kdate3_V, width=2, justify=RIGHT)
        self.kdate1_L = Label(master, text="-")
        self.kdate2_L = Label(master, text="-")
        self.kdate1_E.place(x=470+eltol, y=620)
        self.kdate1_L.place(x=500+eltol, y=620)
        self.kdate2_E.place(x=510+eltol, y=620)
        self.kdate2_L.place(x=525+eltol, y=620)
        self.kdate3_E.place(x=537+eltol, y=620)

        self.soron_kivul_V=BooleanVar()
        self.soron_kivul_B=Checkbutton(master, variable=self.soron_kivul_V, text="Soron kívüli ellátást kért")
        self.soron_kivul_B.place(x=465+eltol, y=640)
        self.elogondozás_V=BooleanVar()
        self.elogondozás_B=Checkbutton(master, variable=self.elogondozás_V, text="Előgondozás")
        self.elogondozás_B.place(x=465+eltol, y=660)
        self.text28=Label(master, text="Dátuma:")
        self.text28.place(x=570+eltol, y=663)
        self.edate1_V = StringVar()
        self.edate2_V = StringVar()
        self.edate3_V = StringVar()
        self.edate1_E = Entry(master, textvariable=self.edate1_V, width=4, justify=RIGHT)
        self.edate2_E = Entry(master, textvariable=self.edate2_V, width=2, justify=RIGHT)
        self.edate3_E = Entry(master, textvariable=self.edate3_V, width=2, justify=RIGHT)
        self.edate1_L = Label(master, text="-")
        self.edate2_L = Label(master, text="-")
        self.edate1_E.place(x=620+eltol, y=663)
        self.edate1_L.place(x=650+eltol, y=663)
        self.edate2_E.place(x=660+eltol, y=663)
        self.edate2_L.place(x=675+eltol, y=663)
        self.edate3_E.place(x=687+eltol, y=663)

        self.ellatas=Label(master, text="Ellátási forma:")
        self.ellatas.place(x=210, y=4)
        self.formak=["Átmeneti Szállás", "Éjjeli Menedékhely", "Nappali Melegedő"]
        self.forma=StringVar()
        self.forma.set(self.formak[0])
        self.temp_forma="Sajtospufi"
        self.formak=tuple(self.formak)
        self.formak_menu = OptionMenu(master, self.forma, *self.formak)
        self.formak_menu.config(width=16)
        self.formak_menu.place(x=210, y=26)
        self.search_for_last()


        self.checker()

    def export3(self):
        dir=askdirectory()
        self.file_header = None
        Msg = messagebox.askyesno("Fejléc betöltése", "Kíván fejlécet rendelni az adatlaphoz?")

        for number in range(0,3):
            self.forma.set(self.formak[number])
            try:
                os.mkdir(dir+"/"+self.formak[number].replace(" ", "_")+"/")
            except:
                pass
            if number==0:
                relative=self.ugyfellista
            elif number==1:
                relative=self.ugyfellista_ejjeli
            else:
                relative=self.ugyfellista_nappali

            for person in relative:
                self.person=str(person)
                self.filename=dir+"/"+self.formak[number].replace(" ", "_")+"/"+self.person+"_"+relative[self.person][0]+"_("+relative[self.person][11]+").docx"
                self.fill_2()
                if Msg==True:
                    self.Open_Header()
                else:
                    self.Open_File()


        self.fill_table()

    def search_for_last(self):
        try:
            if os.path.isfile("auto.csv"):
                self.path="auto.csv"
                #msg=messagebox.showinfo("Biztonsági mentés!", "Haladéktalanul mentse munkáját!")
                self.open_them()
                return


            with open("last.txt", "r") as file:
                self.path=file.readline()
                self.open_them()
        except:
            pass

    def nevjegy(self):
        self.nev=Toplevel()
        self.nev.title("Névjegy")
        root_x = Ablak.winfo_x()
        root_y = Ablak.winfo_y()
        root_width = Ablak.winfo_width()
        root_height = Ablak.winfo_height()
        self.nev.resizable(False,False)
        self.nev.geometry("%dx%d+%d+%d" % (450, 250, root_x + root_width / 2 - 180, root_y + root_height / 2 - 70))
        self.programozo=tkscrolled.ScrolledText(self.nev, width=450, height=250, font=("Times New Roman", 12))
        self.programozo.pack()
        text="Programozta: Fehér János Zoltán\n"
        text2="Verziószám: 1.0\n"
        text3="Frissítés dátumma: 2020.01.30.\n"
        text4="\n"
        text5="Ismerve a szociális szféra lehetőségeit és anyagi forrásait, mivel\njómagam is szociális munkásként tevékenykedtem 10 éven át, a\nszociális ellátórendszer számára minden program, amit a dokumentáció megkönnyítésére készítettem, teljesen ingyenes és szabadon\nfelhasználtható.\n\nKitartás!"
        self.programozo.insert(1.0, text)
        self.programozo.insert(2.0, text2)
        self.programozo.insert(3.0, text3)
        self.programozo.insert(4.0, text4)
        self.programozo.insert(5.0, text5)
        self.programozo.tag_config("bold", font=("Times New Roman", 12, "bold"))
        self.programozo.tag_config("italic", font=("Times New Roman", 12, "italic"))
        self.programozo.tag_add("bold", 1.0, 1.12)
        self.programozo.tag_add("italic", 1.13, 1.45)
        self.programozo.tag_add("bold", 2.0, 2.11)
        self.programozo.tag_add("italic", 2.12, 2.45)
        self.programozo.tag_add("bold", 3.0, 3.18)
        self.programozo.tag_add("italic", 3.19, 3.45)
        self.programozo.config(state=DISABLED)

    def checker(self):
        self.error=False

        regex=re.match("\d\d\d\d-\d\d-\d\d", str(datetime.now())).group()
        self.date=regex.split("-")

        if self.temp_forma!=self.forma.get():
            self.zero_record()


        self.temp_forma=self.forma.get()
        self.active_name = self.list_of_people.get(ACTIVE)
        if self.list_of_people.curselection() != self.now:
            self.fill_table()
        self.now=self.list_of_people.curselection()
        self.check_TAJ()
        self.check_bname()
        self.check_nation()
        self.check_birthdate()
        self.check_phone()
        self.check_gondnok()
        self.iranyitoszam1()
        self.city1()
        self.utcatype1()
        self.iranyitoszam2()
        self.city2()
        self.utcatype2()
        self.egyezik1()
        self.iranyitoszam3()
        self.city3()
        self.utcatype3()
        self.egyezik2()
        self.hazszamok()
        self.check_kdate()
        self.check_edate()
        self.count_auto+=1
        if self.count_auto>250:
            self.path="auto.csv"
            self.save_them()
            self.count_auto=0

        Ablak.after(250, self.checker)







    def check_kdate(self):
        self.kdate1_E.config(bg="white", fg="black")
        self.kdate2_E.config(bg="white", fg="black")
        self.kdate3_E.config(bg="white", fg="black")


        self.kdate1_V.set(self.kdate1_V.get()[:4])
        self.kdate2_V.set(self.kdate2_V.get()[:3])
        self.kdate3_V.set(self.kdate3_V.get()[:3])

        try:
            teszt=int(self.kdate1_V.get())
        except:
            self.kdate1_E.config(bg="red", fg="white")
            self.error = True

        try:
            teszt=int(self.kdate2_V.get())
        except:
            self.kdate2_E.config(bg="red", fg="white")
            self.error = True

        try:
            teszt=int(self.kdate3_V.get())
        except:
            self.kdate3_E.config(bg="red", fg="white")
            self.error = True


        try:

            if int(self.date[0])*365+int(self.date[1])*30+int(self.date[2])<int(self.kdate1_V.get())*365+int(self.kdate2_V.get())*30+int(self.kdate3_V.get()):
                self.kdate1_E.config(bg="red", fg="white")
                self.kdate2_E.config(bg="red", fg="white")
                self.kdate3_E.config(bg="red", fg="white")
                self.error = True


        except:
            pass
        try:
            if int(self.kdate2_V.get()) > 12:
                self.kdate2_E.config(bg="red", fg="white")
                self.error = True
            if int(self.kdate3_V.get())>31:
                self.kdate3_E.config(bg="red", fg="white")
                self.error=True


            elif self.kdate2_V.get() in ["04", "4", "06", "6", "09", "9", "11"] and int(self.kdate3_V.get())>30:
                self.kdate3_E.config(bg="red", fg="white")
                self.error=True

            elif self.kdate2_V.get() in ["02", "2"]:

                if int(self.kdate1_V.get())%4!=0 and int(self.kdate1_V.get())%400!=0 and int(self.kdate3_V.get())>28:
                    self.kdate3_E.config(bg="red", fg="white")
                    self.error = True

                elif self.kdate1_V.get()[2:]=="00" and int(self.kdate1_V.get()[:2])%4!=0 and int(self.kdate3_V.get())>28:
                    self.kdate3_E.config(bg="red", fg="white")
                    self.error = True

                elif int(self.kdate3_V.get()) > 29:
                    self.kdate3_E.config(bg="red", fg="white")
                    self.error = True




        except:
            pass

    def check_edate(self):
        if self.elogondozás_V.get()==1:
            self.edate1_E.config(state=NORMAL)
            self.edate2_E.config(state=NORMAL)
            self.edate3_E.config(state=NORMAL)


            self.edate1_E.config(bg="white", fg="black")
            self.edate2_E.config(bg="white", fg="black")
            self.edate3_E.config(bg="white", fg="black")


            self.edate1_V.set(self.edate1_V.get()[:4])
            self.edate2_V.set(self.edate2_V.get()[:3])
            self.edate3_V.set(self.edate3_V.get()[:3])

            try:
                teszt=int(self.edate1_V.get())
            except:
                self.edate1_E.config(bg="red", fg="white")
                self.error = True

            try:
                if int(self.edate2_V.get())>12:
                    self.edate2_E.config(bg="red", fg="white")
                    self.error=True
                if int(self.edate3_V.get()) > 31:
                    self.edate3_E.config(bg="red", fg="white")
                    self.error = True

                elif self.edate2_V.get() in ["04", "4", "06", "6", "09", "9", "11"] and int(self.edate3_V.get()) > 30:
                    self.edate3_E.config(bg="red", fg="white")
                    self.error = True

                elif self.edate2_V.get() in ["02", "2"]:

                    if int(self.edate1_V.get()) % 4 != 0 and int(self.edate1_V.get()) % 400 != 0 and int(
                            self.edate3_V.get()) > 28:
                        self.edate3_E.config(bg="red", fg="white")
                        self.error = True

                    elif self.edate1_V.get()[2:] == "00" and int(self.edate1_V.get()[:2]) % 4 != 0 and int(
                            self.edate3_V.get()) > 28:
                        self.edate3_E.config(bg="red", fg="white")
                        self.error = True

                    elif int(self.edate3_V.get()) > 29:
                        self.edate3_E.config(bg="red", fg="white")
                        self.error = True



            except:
                pass

            try:
                teszt=int(self.edate2_V.get())
            except:
                self.edate2_E.config(bg="red", fg="white")
                self.error = True

            try:
                teszt=int(self.edate3_V.get())
            except:
                self.edate3_E.config(bg="red", fg="white")
                self.error = True

            try:

                if int(self.date[0])*365+int(self.date[1])*30+int(self.date[2])<int(self.edate1_V.get())*365+int(self.edate2_V.get())*30+int(self.edate3_V.get()):
                    self.edate1_E.config(bg="red", fg="white")
                    self.edate2_E.config(bg="red", fg="white")
                    self.edate3_E.config(bg="red", fg="white")
                    self.error = True


            except:
                pass
        else:
            self.edate1_E.config(state=DISABLED)
            self.edate2_E.config(state=DISABLED)
            self.edate3_E.config(state=DISABLED)

    def hazszamok(self):
        self.address_E.config(bg="white", fg="black")
        self.address2_E.config(bg="white", fg="black")
        self.address3_E.config(bg="white", fg="black")

        try:
            teszt=re.match("\d", self.address_V.get()).group()
        except:
           self.address_E.config(bg="red", fg="white")
           self.error = True


        try:
            teszt=re.match("\d", self.address2_V.get()).group()
        except:
           self.address2_E.config(bg="red", fg="white")
           self.error = True

        try:
            teszt = re.match("\d", self.address3_V.get()).group()
        except:
            self.address3_E.config(bg="red", fg="white")
            self.error = True


    def egyezik1(self):
        if self.megegyezik1.get()==1:
            self.IRSZ_2E.config(state=NORMAL)
            self.city_2E.config(state=NORMAL)

            self.ucca2_E.config(state=NORMAL)
            self.kozterulet_tipusa2_menu.config(state=NORMAL)
            self.kozterulet_tipusa2_egyeb_Entry.config(state=NORMAL)
            self.address2_E.config(state=NORMAL)

            self.IRSZ_2.set(self.IRSZ_1.get())
            self.city_2.set(self.city_1.get().replace("(","").replace(")","").replace(",","").replace("'",""))
            self.ucca2_V.set(self.ucca1_V.get())
            self.kozterulet_tipusa2.set(self.kozterulet_tipusa.get())
            self.kozterulet_tipusa2_egyeb.set(self.kozterulet_tipusa_egyeb.get())
            self.address2_V.set(self.address_V.get())
            self.IRSZ_2E.config(state=DISABLED)
            self.city_2E.config(state=DISABLED)
            self.ucca2_E.config(state=DISABLED)
            self.kozterulet_tipusa2_menu.config(state=DISABLED)
            self.kozterulet_tipusa2_egyeb_Entry.config(state=DISABLED)
            self.address2_E.config(state=DISABLED)

        else:
            self.IRSZ_2E.config(state=NORMAL)
            self.ucca2_E.config(state=NORMAL)
            self.kozterulet_tipusa2_menu.config(state=NORMAL)
            self.address2_E.config(state=NORMAL)


    def egyezik2(self):
        if self.megegyezik2.get()==1:
            self.IRSZ_3E.config(state=NORMAL)
            self.city_3E.config(state=NORMAL)

            self.ucca3_E.config(state=NORMAL)
            self.kozterulet_tipusa3_menu.config(state=NORMAL)
            self.kozterulet_tipusa3_egyeb_Entry.config(state=NORMAL)
            self.address3_E.config(state=NORMAL)

            self.IRSZ_3.set(self.IRSZ_2.get())
            self.city_3.set(self.city_2.get().replace("(","").replace(")","").replace(",","").replace("'",""))
            self.ucca3_V.set(self.ucca2_V.get())
            self.kozterulet_tipusa3.set(self.kozterulet_tipusa2.get())
            self.kozterulet_tipusa3_egyeb.set(self.kozterulet_tipusa2_egyeb.get())
            self.address3_V.set(self.address2_V.get())

            self.IRSZ_3E.config(state=DISABLED)
            self.city_3E.config(state=DISABLED)
            self.ucca3_E.config(state=DISABLED)
            self.kozterulet_tipusa3_menu.config(state=DISABLED)
            self.kozterulet_tipusa3_egyeb_Entry.config(state=DISABLED)
            self.address3_E.config(state=DISABLED)

        else:
            self.IRSZ_3E.config(state=NORMAL)
            self.ucca3_E.config(state=NORMAL)
            self.kozterulet_tipusa3_menu.config(state=NORMAL)
            self.address3_E.config(state=NORMAL)


    def utcatype1(self):
        if self.kozterulet_tipusa.get()=="egyéb:":
            self.kozterulet_tipusa_egyeb_Entry.config(state=NORMAL)
        else:
            self.kozterulet_tipusa_egyeb.set("")
            self.kozterulet_tipusa_egyeb_Entry.config(state=DISABLED)

    def utcatype2(self):
        if self.kozterulet_tipusa2.get()=="egyéb:":
            self.kozterulet_tipusa2_egyeb_Entry.config(state=NORMAL)
        else:
            self.kozterulet_tipusa2_egyeb.set("")
            self.kozterulet_tipusa2_egyeb_Entry.config(state=DISABLED)

    def utcatype3(self):
        if self.kozterulet_tipusa3.get()=="egyéb:":
            self.kozterulet_tipusa3_egyeb_Entry.config(state=NORMAL)
        else:
            self.kozterulet_tipusa3_egyeb.set("")
            self.kozterulet_tipusa3_egyeb_Entry.config(state=DISABLED)

    def iranyitoszam1(self):
        self.IRSZ_1E.config(bg="white", fg="black")
        self.IRSZ_1.set(self.IRSZ_1.get()[:4])

        if len(self.IRSZ_1.get())!=4:
            self.IRSZ_1E.config(bg="red", fg="white")
            self.error = True


        elif self.IRSZ_1.get().startswith("1")==False and self.IRSZ_1.get() not in self.iranyitoszamok:
                self.IRSZ_1E.config(bg="red", fg="white")
                self.error = True

        try:
            teszt=int(self.IRSZ_1.get())
        except:
            self.IRSZ_1E.config(bg="red", fg="white")
            self.error = True


    def iranyitoszam2(self):
        self.IRSZ_2E.config(bg="white", fg="black")
        self.IRSZ_2.set(self.IRSZ_2.get()[:4])

        if len(self.IRSZ_2.get()) != 4:
            self.IRSZ_2E.config(bg="red", fg="white")
            self.error = True


        elif self.IRSZ_2.get().startswith("1") == False and self.IRSZ_2.get() not in self.iranyitoszamok:
            self.IRSZ_2E.config(bg="red", fg="white")
            self.error = True


        try:
            teszt = int(self.IRSZ_2.get())
        except:
            self.IRSZ_2E.config(bg="red", fg="white")
            self.error = True


    def iranyitoszam3(self):
        self.IRSZ_3E.config(bg="white", fg="black")
        self.IRSZ_3.set(self.IRSZ_3.get()[:4])

        if len(self.IRSZ_3.get()) != 4:
            self.IRSZ_3E.config(bg="red", fg="white")
            self.error = True


        elif self.IRSZ_3.get().startswith("1") == False and self.IRSZ_3.get() not in self.iranyitoszamok:
            self.IRSZ_3E.config(bg="red", fg="white")
            self.error = True

        try:
            teszt = int(self.IRSZ_3.get())
        except:
            self.IRSZ_3E.config(bg="red", fg="white")
            self.error = True


    def city1(self):
        self.city_1E.config(state=NORMAL)
        if self.IRSZ_1.get().startswith("1") and len(self.IRSZ_1.get())==4 and self.IRSZ_1.get()[1:3]!="00":
            try:
                if int(self.IRSZ_1.get()[1:3])<23:
                    ker=str("Budapest "+ self.IRSZ_1.get()[1:3]+". kerület")
                    self.city_1.set(ker)
            except:
                pass
        elif self.IRSZ_1.get() in self.iranyitoszamok:

            self.city_1.set(self.iranyitoszamok[self.IRSZ_1.get()][0].replace("(", "").replace( "{", "").replace(")", "").replace( "}", ""))
        self.city_1E.config(state=DISABLED)

    def city2(self):
        self.city_2E.config(state=NORMAL)
        if self.IRSZ_2.get().startswith("1") and len(self.IRSZ_2.get())==4 and self.IRSZ_2.get()[1:3]!="00":
            try:
                if int(self.IRSZ_2.get()[1:3])<23:
                    ker=str("Budapest "+ self.IRSZ_2.get()[1:3]+". kerület")
                    self.city_2.set(ker)
            except:
                pass
        elif self.IRSZ_2.get() in self.iranyitoszamok:
            self.city_2.set(self.iranyitoszamok[self.IRSZ_2.get()][0].replace("(", "").replace( "{", "").replace(")", "").replace( "}", ""))
        self.city_2E.config(state=DISABLED)

    def city3(self):
        self.city_3E.config(state=NORMAL)
        if self.IRSZ_3.get().startswith("1") and len(self.IRSZ_3.get())==4 and self.IRSZ_3.get()[1:3]!="00":
            try:
                if int(self.IRSZ_3.get()[1:3])<23:
                    ker=str("Budapest "+ self.IRSZ_3.get()[1:3]+". kerület")
                    self.city_3.set(ker)
            except:
                pass
        elif self.IRSZ_3.get() in self.iranyitoszamok:
            self.city_3.set(self.iranyitoszamok[self.IRSZ_3.get()][0].replace("(", "").replace( "{", "").replace(")", "").replace( "}", ""))
        self.city_3E.config(state=DISABLED)

    def check_gondnok(self):
        if self.cselekves.get()==1:
            self.gondnok.set(0)
            self.gondnokolt.config(state=DISABLED)

        else:
            self.gondnokolt.config(state=NORMAL)
        if self.gondnok.get()==1:
            self.gondnok_entry.config(state=NORMAL)
            self.gondnok_cim_entry.config(state=NORMAL)
        else:
            self.gondnok_nev.set("")
            self.gondnok_cim.set("")
            self.gondnok_entry.config(state=DISABLED)
            self.gondnok_cim_entry.config(state=DISABLED)


    def check_nation(self):
        if self.nation_var.get()==1:
            self.nationality_var.set("")
            self.idegen.set("")
            self.nationality_entry.config(state=DISABLED)
            self.idegenek.config(state=DISABLED)
        else:
            self.nationality_entry.config(state=NORMAL)
            self.idegenek.config(state=NORMAL)



    def check_bname(self):
        self.bname_entry.config(state=NORMAL)
        if self.same_V.get()==1:
            self.bname_var.set(self.name_var.get())
            self.bname_entry.config(state=DISABLED)

    def check_TAJ(self):
        if self.nincsTAJ.get()==0:
            self.TAJ1_Entry.config(state=NORMAL)
            self.TAJ2_Entry.config(state=NORMAL)
            self.TAJ3_Entry.config(state=NORMAL)
            self.TAJ1_Entry.config(bg="white", fg="black")
            self.TAJ2_Entry.config(bg="white", fg="black")
            self.TAJ3_Entry.config(bg="white", fg="black")
            self.TAJ1_V.set(self.TAJ1_V.get()[:3])
            self.TAJ2_V.set(self.TAJ2_V.get()[:3])
            self.TAJ3_V.set(self.TAJ3_V.get()[:2])

            try:
                teszt=int(self.TAJ1_V.get())
            except:
                self.TAJ1_Entry.config(bg="red", fg="white")
                self.error = True

            try:
                teszt=int(self.TAJ2_V.get())
            except:
                self.TAJ2_Entry.config(bg="red", fg="white")
                self.error = True

            try:
                teszt=int(self.TAJ3_V.get())
            except:
                self.TAJ3_Entry.config(bg="red", fg="white")
                self.error = True

            self.TAJ4_Entry.config(state=NORMAL)
            try:
                self.TAJ4_V.set(str((int(self.TAJ1_V.get()[0])*3+int(self.TAJ1_V.get()[1])*7+int(self.TAJ1_V.get()[2])*3+int(self.TAJ2_V.get()[0])*7+int(self.TAJ2_V.get()[1])*3+int(self.TAJ2_V.get()[2])*7+int(self.TAJ3_V.get()[0])*3++int(self.TAJ3_V.get()[1])*7)%10))
                if self.TAJ4_V.get()=="10":
                    self.TAJ4_V.set('0')
            except:
                self.TAJ4_V.set("")
            self.TAJ4_Entry.config(state=DISABLED)
        else:
            self.TAJ1_V.set("")
            self.TAJ2_V.set("")
            self.TAJ3_V.set("")
            self.TAJ4_V.set("")
            self.TAJ1_Entry.config(state=DISABLED)
            self.TAJ2_Entry.config(state=DISABLED)
            self.TAJ3_Entry.config(state=DISABLED)


    def check_birthdate(self):
        self.bdate1_E.config(bg="white", fg="black")
        self.bdate2_E.config(bg="white", fg="black")
        self.bdate3_E.config(bg="white", fg="black")


        self.bdate1_V.set(self.bdate1_V.get()[:4])
        self.bdate2_V.set(self.bdate2_V.get()[:3])
        self.bdate3_V.set(self.bdate3_V.get()[:3])

        try:
            teszt=int(self.bdate1_V.get())
        except:
            self.bdate1_E.config(bg="red", fg="white")
            self.error = True

        try:
            teszt=int(self.bdate2_V.get())
        except:
            self.bdate2_E.config(bg="red", fg="white")
            self.error = True

        try:
            teszt=int(self.bdate3_V.get())
        except:
            self.bdate3_E.config(bg="red", fg="white")
            self.error = True


        try:
            evek=int(self.date[0])-int(self.bdate1_V.get())
            if int(self.bdate2_V.get())>int(self.date[1]):
                evek-=1
            elif int(self.bdate2_V.get())==self.date[1] and int(self.bdate2_V.get())>int(self.date[2]):
                evek -= 1
            if evek<18:
                self.bdate1_E.config(bg="red", fg="white")
                self.bdate2_E.config(bg="red", fg="white")
                self.bdate3_E.config(bg="red", fg="white")
                self.error = True

        except:
            pass
        try:
            if int(self.bdate2_V.get()) > 12:
                self.bdate2_E.config(bg="red", fg="white")
                self.error = True

            if int(self.bdate3_V.get())>31:
                self.bdate3_E.config(bg="red", fg="white")
                self.error=True

            elif self.bdate2_V.get() in ["04", "4", "06", "6", "09", "9", "11"] and int(self.bdate3_V.get())>30:
                self.bdate3_E.config(bg="red", fg="white")
                self.error=True

            elif self.bdate2_V.get() in ["02", "2"]:

                if int(self.bdate1_V.get())%4!=0 and int(self.bdate1_V.get())%400!=0 and int(self.bdate3_V.get())>28:
                    self.bdate3_E.config(bg="red", fg="white")
                    self.error = True

                elif self.bdate1_V.get()[2:]=="00" and int(self.bdate1_V.get()[:2])%4!=0 and int(self.bdate3_V.get())>28:
                    self.bdate3_E.config(bg="red", fg="white")
                    self.error = True

                elif int(self.bdate3_V.get()) > 29:
                    self.bdate3_E.config(bg="red", fg="white")
                    self.error = True


        except:
            pass

    def check_phone(self):
        self.phone_number1V.config(bg="white", fg="black")
        self.phone_number2V.config(bg="white", fg="black")

        if self.phone_number1.get() not in self.korzetszamok:
            self.phone_number1V.config(bg="red", fg="white")
            self.error = True


        self.phone_number1.set(self.phone_number1.get()[:2])

        if self.phone_number1.get() in ["20", "30", "31", "70"]:
            self.phone_number2.set(self.phone_number2.get()[:7])
            if len(self.phone_number2.get())<7:
                self.phone_number2V.config(bg="red", fg="white")
                self.error = True


        else:
            self.phone_number2.set(self.phone_number2.get()[:6])
            if len(self.phone_number2.get())<6:
                self.phone_number2V.config(bg="red", fg="white")
                self.error = True

        try:
            teszt=int(self.phone_number2.get())
        except:
            self.phone_number2V.config(bg="red", fg="white")
            self.error = True



    def iranyitoszamok_betoltese(self):
        try:
            fieldnames = ["IRSZ", "Település", "Településrész"]
            with open("iranyitoszamok.csv", 'r', newline="") as csvfile:
                next(csvfile)
                reader = csv.DictReader(csvfile, delimiter=';', fieldnames=fieldnames)
                for row in reader:
                    self.iranyitoszamok.setdefault(row["IRSZ"], [])
                    self.iranyitoszamok[row["IRSZ"]].append(row["Település"])
                    if row["Településrész"]!="" and row["Településrész"]!=" " and row["Településrész"]!="None" and row["Településrész"]!=None:
                        try:
                            self.iranyitoszamok[row["IRSZ"]][0]=str(self.iranyitoszamok[row["IRSZ"]][0])+"-"+str(row["Településrész"])
                        except:
                            pass
        except:
            messagebox.showerror("Hiba!", "Nem tölthető be az irányítószám lista!")

    def korzetszamok_betoltese(self):
        try:
            fieldnames = ["Település", "Körzet"]
            with open("korzetszam.csv", 'r', newline="") as csvfile:
                next(csvfile)
                reader = csv.DictReader(csvfile, delimiter=';', fieldnames=fieldnames)
                for row in reader:
                    self.korzetszamok.setdefault(row["Körzet"], [])
                    self.korzetszamok[row["Körzet"]].append(row["Település"])
            self.korzetszamok.setdefault("20", "Telenor - mobil")
            self.korzetszamok.setdefault("30", "Telekom - mobil")
            self.korzetszamok.setdefault("31", "UPC - mobil")
            self.korzetszamok.setdefault("70", "Vodafone - mobil")

        except:
            messagebox.showerror("Hiba!", "Nem tölthető be a körzetszám lista!")

    def new(self):
        if self.changed==True:
            msg=messagebox.askyesnocancel("Nem mentett adatbázis!", "Az ügyféllista megváltozott! Menti a bezárás előtt?")
            if msg=='yes':
                self.save()
            elif msg==None:
                return
        self.zero_record()
        self.list_of_people.select_clear(0, END)
        self.list_of_people.delete(0, END)
        self.ugyfellista={}
        self.ugyfellista_ejjeli = {}
        self.ugyfellista_nappali = {}
        self.person_number.set(1)
        self.person_number_new=1
        self.changed==False
        if os.path.isfile("auto.csv"):
            os.remove("auto.csv")
        self.no_aufo()

    def close(self):
        if self.changed==True:
            msg=messagebox.askyesnocancel("Nem mentett adatbázis!", "Az ügyféllista megváltozott! Menti a bezárás előtt?")
            if msg=='yes':
                self.save()
            elif msg==None:
                return
        Ablak.destroy()

    def open(self):
        self.path = askopenfilename(initialdir="*", title="Fájl kiválasztása", filetypes=(("csv fájlok", "*.csv"),))
        self.open_them()

    def open_them(self):
        if self.changed==True:
            msg=messagebox.askyesno("Nem mentett adatbázis!", "Az ügyféllista megváltozott! Menti a bezárás előtt?")
            if msg=='yes':
                self.save()

        self.ugyfellista={}
        self.ugyfellista_ejjeli={}
        self.ugyfellista_nappali={}
        self.list_of_people.select_clear(0, END)
        self.list_of_people.delete(0, END)
        fieldnames = ["Ellatas", "ID"]
        for number in range(0, 52):
            fieldnames.append(str(number))
        with open(self.path, 'r', newline="") as csvfile:

            reader = csv.DictReader(csvfile, delimiter=';', fieldnames=fieldnames)
            reader.__next__()
            for row in reader:
                if row["Ellatas"]=="HÁO":
                    self.ugyfellista.setdefault(row["ID"],[])
                    for number in range(0, 52):
                        self.ugyfellista[row["ID"]].append(row[str(number)])
                    self.list_of_people.insert(END, row["ID"]+".: "+row["0"]+" ("+row["11"]+")")
                elif row["Ellatas"] == "ÉM":
                    self.ugyfellista_ejjeli.setdefault(row["ID"], [])
                    for number in range(0, 52):
                        self.ugyfellista_ejjeli[row["ID"]].append(row[str(number)])

                else:
                    self.ugyfellista_nappali.setdefault(row["ID"], [])
                    for number in range(0, 52):
                        self.ugyfellista_nappali[row["ID"]].append(row[str(number)])

            for ugyfel in self.ugyfellista:
                if len(str(self.ugyfellista[ugyfel][1])) < 3:
                    self.ugyfellista[ugyfel][1] = str("0" * (3 - len(str(self.ugyfellista[ugyfel][1]))) + str(self.ugyfellista[ugyfel][1]))
                if len(str(self.ugyfellista[ugyfel][2])) < 3:
                    self.ugyfellista[ugyfel][2] = str("0" * (3 - len(str(self.ugyfellista[ugyfel][2]))) + str(self.ugyfellista[ugyfel][2]))
                if len(str(self.ugyfellista[ugyfel][3])) < 2:
                    self.ugyfellista[ugyfel][3] = str("0" * (2 - len(str(self.ugyfellista[ugyfel][3]))) + str(self.ugyfellista[ugyfel][3]))
                if len(str(self.ugyfellista[ugyfel][12])) < 2:
                    self.ugyfellista[ugyfel][12] = str("0" * (2 - len(str(self.ugyfellista[ugyfel][12]))) + str(self.ugyfellista[ugyfel][12]))
                if len(str(self.ugyfellista[ugyfel][13])) < 2:
                    self.ugyfellista[ugyfel][13] = str("0" * (2 - len(str(self.ugyfellista[ugyfel][13]))) + str(self.ugyfellista[ugyfel][13]))

            for ugyfel in self.ugyfellista_ejjeli:
                if len(str(self.ugyfellista_ejjeli[ugyfel][1])) < 3:
                    self.ugyfellista_ejjeli[ugyfel][1] = str("0" * (3 - len(str(self.ugyfellista_ejjeli[ugyfel][1]))) + str(self.ugyfellista_ejjeli[ugyfel][1]))
                if len(str(self.ugyfellista_ejjeli[ugyfel][2])) < 3:
                    self.ugyfellista_ejjeli[ugyfel][2] = str("0" * (3 - len(str(self.ugyfellista_ejjeli[ugyfel][2]))) + str(self.ugyfellista_ejjeli[ugyfel][2]))
                if len(str(self.ugyfellista_ejjeli[ugyfel][3])) < 2:
                    self.ugyfellista_ejjeli[ugyfel][3] = str("0" * (2 - len(str(self.ugyfellista_ejjeli[ugyfel][3]))) + str(self.ugyfellista_ejjeli[ugyfel][3]))
                if len(str(self.ugyfellista_ejjeli[ugyfel][12])) < 2:
                    self.ugyfellista_ejjeli[ugyfel][12] = str("0" * (2 - len(str(self.ugyfellista_ejjeli[ugyfel][12]))) + str(self.ugyfellista_ejjeli[ugyfel][12]))
                if len(str(self.ugyfellista_ejjeli[ugyfel][13])) < 2:
                    self.ugyfellista_ejjeli[ugyfel][13] = str("0" * (2 - len(str(self.ugyfellista_ejjeli[ugyfel][13]))) + str(self.ugyfellista_ejjeli[ugyfel][13]))

            for ugyfel in self.ugyfellista_nappali:
                if len(str(self.ugyfellista_nappali[ugyfel][1])) < 3:
                    self.ugyfellista_nappali[ugyfel][1] = str("0" * (3 - len(str(self.ugyfellista_nappali[ugyfel][1]))) + str(self.ugyfellista_nappali[ugyfel][1]))
                if len(str(self.ugyfellista_nappali[ugyfel][2])) < 3:
                    self.ugyfellista_nappali[ugyfel][2] = str("0" * (3 - len(str(self.ugyfellista_nappali[ugyfel][2]))) + str(self.ugyfellista_nappali[ugyfel][2]))
                if len(str(self.ugyfellista_nappali[ugyfel][3])) < 2:
                    self.ugyfellista_nappali[ugyfel][3] = str("0" * (2 - len(str(self.ugyfellista_nappali[ugyfel][3]))) + str(self.ugyfellista_nappali[ugyfel][3]))
                if len(str(self.ugyfellista_nappali[ugyfel][12])) < 2:
                    self.ugyfellista_nappali[ugyfel][12] = str("0" * (2 - len(str(self.ugyfellista_nappali[ugyfel][12]))) + str(self.ugyfellista_nappali[ugyfel][12]))
                if len(str(self.ugyfellista_nappali[ugyfel][13])) < 2:
                    self.ugyfellista_nappali[ugyfel][13] = str("0" * (2 - len(str(self.ugyfellista_nappali[ugyfel][13]))) + str(self.ugyfellista_nappali[ugyfel][13]))

            self.forma.set(self.formak[0])

            self.list_of_people.see(END)
            self.zero_record()
            self.no_aufo()

    def save(self):
        if self.path=="":
            self.save_as()
            return
        if os.path.isfile("auto.csv"):
            os.remove("auto.csv")
        self.save_them()

    def save_as(self):
        file = asksaveasfile(initialdir="*", title="Fájl kiválasztása", filetypes=(("csv fájlok", "*.csv"),))
        file.close()
        self.path = file.name
        if os.path.isfile("auto.csv"):
            os.remove("auto.csv")
        self.save_them()

    def save_them(self):
        try:
            os.remove(self.path)
        except:
            pass
        fieldnames=["Ellatas", "ID"]
        for number in range(0,52):
            fieldnames.append(str(number))

        if self.path.endswith(".csv") is False:
            self.path = self.path + ".csv"
        with open(self.path, 'w', newline="") as csvfile:
            csv_writer = csv.DictWriter(csvfile, delimiter=';', fieldnames=fieldnames)
            csv_writer.writeheader()
            try:
                for key in self.ugyfellista:
                    templista={}
                    templista['Ellatas']="HÁO"
                    templista['ID']=key
                    for number in range(0,52):
                        templista[str(number)]=self.ugyfellista[key][number]
                    csv_writer.writerow(templista)
            except:
                pass
            try:
                for key in self.ugyfellista_ejjeli:
                    templista={}
                    templista['Ellatas'] = "ÉM"
                    templista['ID']=key
                    for number in range(0,52):
                        templista[str(number)]=self.ugyfellista_ejjeli[key][number]
                    csv_writer.writerow(templista)
            except:
                pass
            try:
                for key in self.ugyfellista_nappali:
                    templista = {}
                    templista['Ellatas'] = "NM"
                    templista['ID'] = key
                    for number in range(0, 52):
                        templista[str(number)] = self.ugyfellista_nappali[key][number]
                    csv_writer.writerow(templista)
            except:
                pass
        if self.path!="auto.csv":
            with open("last.txt", "w") as file:
                file.write(self.path)
        self.no_aufo()
        self.changed = False

    def no_aufo(self):
        try:
            if os.path.isfile("auto.csv"):
                os.remove("auto.csv")
        except:
            pass

    def export1(self):
        try:
            file = asksaveasfile(initialdir="*", title="Fájl kiválasztása", filetypes=(("excel 2007 fájlok", "*.xlsx"),))

            path = file.name
            file.close()
            os.remove(path)
            if path.endswith(".xlsx"):
                pass
            else:
                path = path + ".xlsx"
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = "Átmeneti Szállás"
            worksheet_ejjeli=workbook.create_sheet(title="Éjjeli Menedékhely")
            worksheet_nappali=workbook.create_sheet(title="Nappali Melegedő")


            for number in range(0,3):
                workbook.active=number
                worksheet=workbook.active
                if number==0:
                    relative=self.ugyfellista
                elif number==1:
                    relative=self.ugyfellista_ejjeli
                else:
                    relative=self.ugyfellista_nappali

                worksheet = workbook.active
                worksheet["A1"] = "ID"
                worksheet["B1"] = "TAJ-szám"
                worksheet["C1"] = "Név"
                worksheet["D1"] = "Születési Név"
                worksheet["E1"] = "Születési Hely"
                worksheet["F1"] = "Születési Idő"
                worksheet["G1"] = "Édesanyja neve"
                worksheet["H1"] = "Állampolgárság"
                worksheet["I1"] = "Idegenrendészeti státusz"
                worksheet["J1"] = "Cselekvőképesség"
                worksheet["K1"] = "Gondnokolt"
                worksheet["L1"] = "Gondnok neve"
                worksheet["M1"] = "Gondnok elérhetősége"

                worksheet.merge_cells("A1:A2")
                worksheet.merge_cells("B1:B2")
                worksheet.merge_cells("C1:C2")
                worksheet.merge_cells("D1:D2")
                worksheet.merge_cells("E1:E2")
                worksheet.merge_cells("F1:F2")
                worksheet.merge_cells("G1:G2")
                worksheet.merge_cells("H1:H2")
                worksheet.merge_cells("I1:I2")
                worksheet.merge_cells("J1:J2")
                worksheet.merge_cells("K1:K2")
                worksheet.merge_cells("L1:L2")
                worksheet.merge_cells("M1:M2")

                worksheet["N1"] = "Utolsó állandó lakhely"
                worksheet["N2"] = "Irányítószám"
                worksheet["O2"] = "Város"
                worksheet["P2"] = "Közterület neve"
                worksheet["Q2"] = "Közterület típusa"
                worksheet["R2"] = "Házszám"
                worksheet.merge_cells("N1:R1")

                worksheet["S1"] = "Tartózkodási hely"
                worksheet["S2"] = "Irányítószám"
                worksheet["T2"] = "Város"
                worksheet["U2"] = "Közterület neve"
                worksheet["V2"] = "Közterület típusa"
                worksheet["W2"] = "Házszám"
                worksheet.merge_cells("S1:W1")

                worksheet["X1"] = "Értesítési cím"
                worksheet["X2"] = "Irányítószám"
                worksheet["Y2"] = "Város"
                worksheet["Z2"] = "Közterület neve"
                worksheet["AA2"] = "Közterület típusa"
                worksheet["AB2"] = "Házszám"
                worksheet.merge_cells("X1:AB1")



                worksheet["AC1"] = "Megnevezett nagykorú hozzátartozó"
                worksheet["AD1"] = "Hozzátartozó elérhetősége"
                worksheet["AE1"] = "Kérelem beterjesztése:"
                worksheet["AF1"] = "Soron kívüli ellátást kért?"
                worksheet["AG1"] = "Előgondozás történt?"
                worksheet["AH1"] = "Előgondozás dátuma"
                worksheet["AI1"] = "Kérelmező telefonszáma"


                worksheet.merge_cells("AC1:AC2")
                worksheet.merge_cells("AD1:AD2")
                worksheet.merge_cells("AE1:AE2")
                worksheet.merge_cells("AF1:AF2")
                worksheet.merge_cells("AG1:AG2")
                worksheet.merge_cells("AH1:AH2")
                worksheet.merge_cells("AI1:AI2")



                for cell in worksheet["1"]:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.font = Font(bold=True)
                for cell in worksheet["2"]:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.font = Font(bold=True)
                counter=2
                for ugyfel in relative:
                    counter += 1
                    worksheet["A" + str(counter)] = ugyfel
                    worksheet["A" + str(counter)].alignment = Alignment(horizontal='right', vertical='center')
                    if relative[ugyfel][1] == "" or relative[ugyfel][2] == "" or relative[ugyfel][3] == "":
                        worksheet["B" + str(counter)] = "Hiányzó TAJ szám!!"
                        worksheet["B" + str(counter)].font = Font(bold=True, color=RED)
                    else:
                        worksheet["B" + str(counter)] = relative[ugyfel][1] + "-" + relative[ugyfel][2] + "-" + \
                                                        relative[ugyfel][3] + relative[ugyfel][4]

                    worksheet["C" + str(counter)] = relative[ugyfel][0]
                    worksheet["C" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["D" + str(counter)] = relative[ugyfel][6]
                    worksheet["D" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["E" + str(counter)] = relative[ugyfel][10]
                    worksheet["E" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    if relative[ugyfel][11] == "" or relative[ugyfel][12] == "" or relative[ugyfel][13] == "":
                        worksheet["F" + str(counter)] = "Nincs megadva születési dátum!"
                        worksheet["F" + str(counter)].font = Font(bold=True, color=RED)
                        worksheet["F" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    else:
                        worksheet["F" + str(counter)] = relative[ugyfel][11] + "." + relative[ugyfel][12] + "." + \
                                                        relative[ugyfel][13] + "."
                        try:
                            worksheet["F" + str(counter)].style = NamedStyle(name='datetime', number_format='YYYY/MM/DD')
                        except:
                            worksheet["F" + str(counter)].style = "datetime"

                    worksheet["F" + str(counter)].alignment = Alignment(horizontal='right', vertical='center')

                    worksheet["G" + str(counter)] = relative[ugyfel][49]
                    worksheet["G" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')

                    if relative[ugyfel][8] == "1":
                        worksheet["H" + str(counter)] = "Magyar"
                    else:
                        worksheet["H" + str(counter)] = relative[ugyfel][9]
                    worksheet["H" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    worksheet["I" + str(counter)] = relative[ugyfel][16]
                    worksheet["I" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    if relative[ugyfel][17] == "1":
                        worksheet["J" + str(counter)] = "Cselekvőképes"
                    elif relative[ugyfel][17] == "2":
                        worksheet["J" + str(counter)] = "Korlátozottan Cselekvőképes"

                    elif relative[ugyfel][17] == "3":
                        worksheet["J" + str(counter)] = "Cselekvőképtelen"

                    worksheet["J" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    if relative[ugyfel][18] == 1 or relative[ugyfel][19] == True or relative[ugyfel][20] == "1":
                        worksheet["K" + str(counter)] = "Igen"
                        worksheet["L" + str(counter)] = relative[ugyfel][50]
                        worksheet["M" + str(counter)] = relative[ugyfel][51]
                        worksheet["L" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                        worksheet["M" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    else:
                        worksheet["K" + str(counter)] = "Nem"
                    worksheet["K" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    worksheet["N" + str(counter)] = relative[ugyfel][19]
                    worksheet["N" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    worksheet["O" + str(counter)] = relative[ugyfel][20].replace("('", "").replace("',)", "")
                    worksheet["O" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["P" + str(counter)] = relative[ugyfel][21]
                    worksheet["P" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')

                    if relative[ugyfel][22] == "egyéb:":
                        worksheet["Q" + str(counter)] = relative[ugyfel][23]
                    else:
                        worksheet["Q" + str(counter)] = relative[ugyfel][22]
                    worksheet["Q" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["R" + str(counter)] = relative[ugyfel][24]
                    worksheet["R" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    worksheet["S" + str(counter)] = relative[ugyfel][25]
                    worksheet["S" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    worksheet["T" + str(counter)] = relative[ugyfel][26]
                    worksheet["T" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["U" + str(counter)] = relative[ugyfel][27]
                    worksheet["U" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    if relative[ugyfel][27] == "egyéb:":
                        worksheet["V" + str(counter)] = relative[ugyfel][29]
                    else:
                        worksheet["V" + str(counter)] = relative[ugyfel][28]
                    worksheet["V" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["W" + str(counter)] = relative[ugyfel][30]
                    worksheet["W" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    worksheet["X" + str(counter)] = relative[ugyfel][31]
                    worksheet["X" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    worksheet["Y" + str(counter)] = relative[ugyfel][32]
                    worksheet["Y" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["Z" + str(counter)] = relative[ugyfel][33]
                    worksheet["Z" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    if relative[ugyfel][34] == "egyéb:":
                        worksheet["AA" + str(counter)] = relative[ugyfel][35]
                    else:
                        worksheet["AA" + str(counter)] = relative[ugyfel][34]
                    worksheet["AA" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')
                    worksheet["AB" + str(counter)] = relative[ugyfel][36]
                    worksheet["AB" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    worksheet["AC" + str(counter)] = relative[ugyfel][39]
                    worksheet["AC" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    worksheet["AD" + str(counter)] = relative[ugyfel][40]
                    worksheet["AD" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')
                    if relative[ugyfel][41] == 1 or relative[ugyfel][41] == True or relative[ugyfel][41] == "1" or \
                            relative[ugyfel][41] == "True":
                        worksheet["AF" + str(counter)] = "Igen"
                    else:
                        worksheet["AF" + str(counter)] = "Nem"
                    worksheet["AF" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    if relative[ugyfel][43] == "" or relative[ugyfel][44] == "" or relative[ugyfel][
                        45] == "":
                        worksheet["AE" + str(counter)] = "Nincs megadva előterjesztési dátum!!"
                        worksheet["AE" + str(counter)].font = Font(bold=True, color=RED)
                        worksheet["AE" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    else:

                        worksheet["AE" + str(counter)] = relative[ugyfel][43] + "." + relative[ugyfel][
                            44] + "." + relative[ugyfel][45] + "."
                        try:
                            worksheet["AE" + str(counter)].style = NamedStyle(name='datetime',
                                                                              number_format='YYYY/MM/DD')
                        except:
                            worksheet["AE" + str(counter)].style = "datetime"
                            worksheet["AE" + str(counter)].alignment = Alignment(horizontal='right', vertical='center')

                    if relative[ugyfel][42] == True or relative[ugyfel][42] == "True" or relative[ugyfel][42] == 1 or \
                            relative[ugyfel][42] == "1":
                        worksheet["AG" + str(counter)] = "Igen"
                        if relative[ugyfel][46] == "" or relative[ugyfel][47] == "" or \
                                relative[ugyfel][48] == "":

                            worksheet["AH" + str(counter)] = "Nincs megadva előgondozási dátum!!"
                            worksheet["AH" + str(counter)].font = Font(bold=True, color=RED)
                            worksheet["AH" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                        else:
                            worksheet["AH" + str(counter)] = relative[ugyfel][46] + "." + relative[ugyfel][
                                47] + "." + relative[ugyfel][48] + "."
                            try:
                                worksheet["AH" + str(counter)].style = NamedStyle(name='datetime',
                                                                                  number_format='YYYY/MM/DD')
                            except:
                                worksheet["AH" + str(counter)].style = "datetime"
                                worksheet["AH" + str(counter)].alignment = Alignment(horizontal='right', vertical='center')


                    else:
                        worksheet["AG" + str(counter)] = "Nem"
                    worksheet["AG" + str(counter)].alignment = Alignment(horizontal='center', vertical='center')

                    if relative[ugyfel][14] != "" and relative[ugyfel][15] != "":
                        worksheet["AI" + str(counter)] = str(relative[ugyfel][14]) + "/" + str(
                            relative[ugyfel][15][:3] + "-" + str(relative[ugyfel][15][3:]))
                        worksheet["AI" + str(counter)].alignment = Alignment(horizontal='left', vertical='center')

                outside_border_left_up = Border(
                    left=Side(border_style="double", color='00000000'),
                    right=Side(border_style=None, color='00000000'),
                    top=Side(border_style="double", color='00000000'),
                    bottom=Side(border_style="thin", color='00000000')
                )
                outside_border_right_up = Border(
                    left=Side(border_style=None, color='00000000'),
                    right=Side(border_style="double", color='00000000'),
                    top=Side(border_style="double", color='00000000'),
                    bottom=Side(border_style="thin", color='00000000')
                )
                outside_border_up = Border(
                    left=Side(border_style=None, color='00000000'),
                    right=Side(border_style=None, color='00000000'),
                    top=Side(border_style="double", color='00000000'),
                    bottom=Side(border_style="thin", color='00000000')
                )

                outside_border_left = Border(
                    left=Side(border_style="double", color='00000000'),
                    right=Side(border_style=None, color='00000000'),
                    top=Side(border_style="thin", color='00000000'),
                    bottom=Side(border_style="thin", color='00000000')
                )

                outside_border_right = Border(
                    left=Side(border_style=None, color='00000000'),
                    right=Side(border_style="double", color='00000000'),
                    top=Side(border_style="thin", color='00000000'),
                    bottom=Side(border_style="thin", color='00000000')
                )

                outside_border_middle = Border(
                    left=Side(border_style=None, color='00000000'),
                    right=Side(border_style=None, color='00000000'),
                    top=Side(border_style="thin", color='00000000'),
                    bottom=Side(border_style="thin", color='00000000')
                )

                outside_border_left_down = Border(
                    left=Side(border_style="double", color='00000000'),
                    right=Side(border_style=None, color='00000000'),
                    top=Side(border_style="thin", color='00000000'),
                    bottom=Side(border_style="double", color='00000000')
                )

                outside_border_right_down = Border(
                    left=Side(border_style=None, color='00000000'),
                    right=Side(border_style="double", color='00000000'),
                    top=Side(border_style="thin", color='00000000'),
                    bottom=Side(border_style="double", color='00000000')
                )

                outside_border_down = Border(
                    left=Side(border_style=None, color='00000000'),
                    right=Side(border_style=None, color='00000000'),
                    top=Side(border_style="thin", color='00000000'),
                    bottom=Side(border_style="double", color='00000000')
                )
                column_list = []
                for column in worksheet.columns:
                    column_list.append(column)

                for column in column_list:
                    for cell in column:
                        if cell == column[0]:
                            if column == column_list[0]:
                                cell.border = outside_border_left_up
                            elif column == column_list[-1]:
                                cell.border = outside_border_right_up
                            else:
                                cell.border = outside_border_up
                        elif cell == column[-1]:
                            if column == column_list[0]:
                                cell.border = outside_border_left_down
                            elif column == column_list[-1]:
                                cell.border = outside_border_right_down
                            else:
                                cell.border = outside_border_down
                        else:
                            if column == column_list[0]:
                                cell.border = outside_border_left
                            elif column == column_list[-1]:
                                cell.border = outside_border_right
                            else:
                                cell.border = outside_border_middle

                for column in worksheet.columns:
                    oszlop_hossz = 0
                    for cell in column:
                        if cell.value != None:
                            try:
                                if len(cell.value) + 5 > oszlop_hossz:
                                    oszlop_hossz = len(cell.value) + 5
                            except:
                                if 6 > oszlop_hossz:
                                    oszlop_hossz = 6

                            worksheet.column_dimensions[cell.column].width = oszlop_hossz

                workbook.save(path)

        except:
            Msg = messagebox.showerror("Hiba!", "Az adatbázis Excel fájlként mentése sikertelen!")



    def export2(self):
      self.file_header = None

      try:
        if self.list_of_people.curselection()=="":
            Msg=messagebox.showerror("Hiba!", "Nincs ügyfél kijelölve!")
        else:
            document = Document()
            file = asksaveasfile(initialdir="*", title="Fájl kiválasztása", filetypes=(("Word 2007 fájlok", "*.docx"),))
            self.filename = file.name
            file.close()
            os.remove(self.filename)
            if self.filename.endswith(".docx") is False:
                self.filename = self.filename + ".docx"

            Msg=messagebox.askyesno("Fejléc betöltése", "Kíván fejlécet rendelni az adatlaphoz?")
            if Msg==True:
                self.Open_Header()
            else:
                self.Open_File()
      except:
          msg = messagebox.showerror("Hiba!", "Fejléc importálása sikertelen!")

    def Open_Header(self):
        try:
            if self.file_header==None:
                self.file_header=askopenfilename(initialdir = "*",title = "Fájl kiválasztása", filetypes = (("Word 2007 fájlok","*.docx"),))
            header_doc=Document(self.file_header)

            file=open(self.filename, "w")
            self.save_doc=header_doc


            self.save_doc._body.clear_content()

            self.Write_File()
        except:
            msg=messagebox.showerror("Hiba!", "Fejléc importálása sikertelen!")


    def Open_File(self):
        try:
            file=open(self.filename, "w")
            self.save_doc = Document()
            self.Write_File()
        except:
            msg=messagebox.showerror("Hiba!", "Adatlap exportálása sikertelen!")

    def Write_File(self):
            sorszam_para=self.save_doc.add_paragraph()
            sorszam_datum=sorszam_para.add_run("Kérelem dátuma:")
            sorszam_datum.font.bold=True
            sorszam_datum.font.underline=True
            sorszam_datum.font.italic=False
            sorszam_datum.font.size=Pt(12)
            sorszam_datum.font.name="Times New Roman"
            sorszam_datum2=sorszam_para.add_run("\t")
            sorszam_datum2.font.bold=True
            sorszam_datum2.font.underline=False
            sorszam_datum2.font.italic=False
            sorszam_datum2.font.size=Pt(12)
            sorszam_datum2.font.name="Times New Roman"
            for_date=self.kdate1_V.get()+"."+self.kdate2_V.get()+"."+self.kdate3_V.get()+"."
            try:
                regex=re.match("\d\d\d\d\.\d\d\.\d\d\.", for_date).group()
                sorszam_text3 = sorszam_para.add_run(for_date)
            except:
                sorszam_text3 = sorszam_para.add_run("________.____.____.")
            sorszam_text3.font.bold = True
            sorszam_text3.font.underline = False
            sorszam_text3.font.italic= True
            sorszam_text3.font.size = Pt(12)
            sorszam_text3.font.name = "Times New Roman"
            sorszam_datum3 = sorszam_para.add_run("\t")
            sorszam_datum3.font.bold = True
            sorszam_datum3.font.underline = False
            sorszam_datum3.font.italic = False
            sorszam_datum3.font.size = Pt(12)
            sorszam_datum3.font.name = "Times New Roman"

            sorszam_text1 = sorszam_para.add_run("Ügyfél sorszáma:")
            sorszam_text1.font.bold=True
            sorszam_text1.font.underline=True
            sorszam_text1.font.size=Pt(12)
            sorszam_text1.font.name="Times New Roman"
            sorszam_text2=sorszam_para.add_run("\t"+str(self.person_number.get()))
            sorszam_text2.font.bold=True
            sorszam_text2.font.italic=True
            sorszam_text2.font.underline=False
            sorszam_text2.font.name = "Times New Roman"
            sorszam_text2.font.size=Pt(12)
            sorszam_para.paragraph_format.space_before = Pt(5)
            sorszam_para.paragraph_format.space_after = Pt(15)
            sorszam_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            sorszam_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            sorszam_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            sorszam_para.paragraph_format.line_spacing=1


            title_para=self.save_doc.add_paragraph()
            title_text1=title_para.add_run("Nyilvántartó Adatlap")
            title_text1.font.bold=True
            title_text1.font.italic=False
            title_text1.font.underline=True
            title_text1.font.size=Pt(22)
            title_text1.font.name = "Times New Roman"
            title_para.paragraph_format.space_before = Pt(20)
            title_para.paragraph_format.space_after = Pt(35)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.line_spacing=1

            ellatas_tipusa = self.save_doc.add_paragraph()
            ellatas_text = ellatas_tipusa.add_run("Ellátás típusa:")
            ellatas_text.font.bold = True
            ellatas_text.font.italic = False
            ellatas_text.font.underline = True
            ellatas_text.font.size = Pt(12)
            ellatas_text.font.name = "Times New Roman"
            ellatas_text3 = ellatas_tipusa.add_run("\t")
            ellatas_text3.font.bold = True
            ellatas_text3.font.italic = False
            ellatas_text3.font.underline = False
            ellatas_text3.font.size = Pt(12)
            ellatas_text3.font.name = "Times New Roman"


            ellatas_text2 = ellatas_tipusa.add_run(self.forma.get())
            ellatas_text2.font.bold = True
            ellatas_text2.font.italic = True
            ellatas_text2.font.underline = False
            ellatas_text2.font.size = Pt(12)
            ellatas_text2.font.name = "Times New Roman"

            ellatas_tipusa.paragraph_format.space_before = Pt(40)
            ellatas_tipusa.paragraph_format.space_after = Pt(0)
            ellatas_tipusa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


            ellatas_tipusa.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


            nev_para=self.save_doc.add_paragraph()
            nev_text1=nev_para.add_run("Ügyfél neve:")
            nev_text1.font.bold=True
            nev_text1.font.italic=False
            nev_text1.font.underline=True
            nev_text1.font.size=Pt(12)
            nev_text1.font.name = "Times New Roman"
            nev_text2=nev_para.add_run("\t")
            nev_text2.font.bold=True
            nev_text2.font.italic=False
            nev_text2.font.underline=False
            nev_text2.font.size=Pt(12)
            nev_text2.font.name = "Times New Roman"
            nev_text3=nev_para.add_run(self.name_var.get())
            nev_text3.font.bold=True
            nev_text3.font.italic=True
            nev_text3.font.underline=False
            nev_text3.font.size=Pt(12)
            nev_text3.font.name = "Times New Roman"



            taj_text0=nev_para.add_run("\t")
            taj_text0.font.bold = True
            taj_text0.font.italic = False
            taj_text0.font.underline = False
            taj_text0.font.size = Pt(12)
            taj_text0.font.name = "Times New Roman"
            taj_text1=nev_para.add_run("TAJ-szám.:")
            taj_text1.font.bold = True
            taj_text1.font.italic = False
            taj_text1.font.underline = True
            taj_text1.font.size = Pt(12)
            taj_text1.font.name = "Times New Roman"
            taj_text2 = nev_para.add_run("\t")
            taj_text2.font.bold = True
            taj_text2.font.italic = False
            taj_text2.font.underline = False
            taj_text2.font.size = Pt(12)
            taj_text2.font.name = "Times New Roman"

            for_taj=self.TAJ1_V.get()+"-"+self.TAJ2_V.get()+"-"+self.TAJ3_V.get()+self.TAJ4_V.get()
            try:
                regex=re.match("\d\d\d-\d\d\d-\d\d\d", for_taj).group()
                taj_text3 = nev_para.add_run(" "+for_taj)

            except:
                taj_text3 = nev_para.add_run(" _____-_____-_____")

            taj_text3.font.bold = True
            taj_text3.font.italic = True
            taj_text3.font.underline = False
            taj_text3.font.size = Pt(12)
            taj_text3.font.name = "Times New Roman"

            nev_para.paragraph_format.space_before = Pt(0)
            nev_para.paragraph_format.space_after = Pt(0)
            nev_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            nev_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            nev_para.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            nev_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            nev_para.paragraph_format.line_spacing = 1.25

            szul_nev_para=self.save_doc.add_paragraph()
            szul_nev_text1=szul_nev_para.add_run("Születési név:")
            szul_nev_text1.font.bold = True
            szul_nev_text1.font.italic = False
            szul_nev_text1.font.underline = True
            szul_nev_text1.font.size = Pt(12)
            szul_nev_text1.font.name = "Times New Roman"
            szul_nev_text2=szul_nev_para.add_run("\t")
            szul_nev_text2.font.bold = True
            szul_nev_text2.font.italic = False
            szul_nev_text2.font.underline = False
            szul_nev_text2.font.size = Pt(12)
            szul_nev_text2.font.name = "Times New Roman"
            szul_nev_text3=szul_nev_para.add_run(self.bname_var.get())
            szul_nev_text3.font.bold = True
            szul_nev_text3.font.italic = True
            szul_nev_text3.font.underline = False
            szul_nev_text3.font.size = Pt(12)
            szul_nev_text3.font.name = "Times New Roman"

            anyja_neve0=szul_nev_para.add_run("\t")
            anyja_neve0.font.bold = True
            anyja_neve0.font.italic = False
            anyja_neve0.font.underline = False
            anyja_neve0.font.name = "Times New Roman"
            anyja_neve0.font.size = Pt(12)

            anyja_neve=szul_nev_para.add_run("Anyja neve:")
            anyja_neve.font.bold=True
            anyja_neve.font.italic=False
            anyja_neve.font.underline=True
            anyja_neve.font.name="Times New Roman"
            anyja_neve.font.size= Pt(12)
            anyja_neve3=szul_nev_para.add_run("\t")
            anyja_neve3.font.bold = True
            anyja_neve3.font.italic = False
            anyja_neve3.font.underline = False
            anyja_neve3.font.name = "Times New Roman"
            anyja_neve3.font.size = Pt(12)

            anyja_neve2=szul_nev_para.add_run(self.mothername.get())
            anyja_neve2.font.bold = True
            anyja_neve2.font.italic = True
            anyja_neve2.font.underline = False
            anyja_neve2.font.name = "Times New Roman"
            anyja_neve2.font.size = Pt(12)


            szul_nev_para.paragraph_format.space_before = Pt(0)
            szul_nev_para.paragraph_format.space_after = Pt(0)
            szul_nev_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            szul_nev_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            szul_nev_para.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            szul_nev_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)



            szul_hely_para=self.save_doc.add_paragraph()
            szul_hely_text1=szul_hely_para.add_run("Születési hely:")
            szul_hely_text1.font.bold=True
            szul_hely_text1.font.italic=False
            szul_hely_text1.font.size=Pt(12)
            szul_hely_text1.font.name = "Times New Roman"


            szul_hely_text1.font.underline=True
            szul_hely_text2=szul_hely_para.add_run("\t")
            szul_hely_text2.font.bold=True
            szul_hely_text2.font.italic=False
            szul_hely_text2.font.underline=False
            szul_hely_text2.font.size=Pt(12)
            szul_hely_text2.font.name = "Times New Roman"

            szul_hely_text3=szul_hely_para.add_run(self.bplace_V.get())
            szul_hely_text3.font.bold=True
            szul_hely_text3.font.italic=True
            szul_hely_text3.font.underline=False
            szul_hely_text3.font.size=Pt(12)
            szul_hely_text3.font.name = "Times New Roman"


            szul_hely_text4=szul_hely_para.add_run("\t")
            szul_hely_text4.font.bold=True
            szul_hely_text4.font.italic=False
            szul_hely_text4.font.underline=False
            szul_hely_text4.font.size=Pt(12)
            szul_hely_text4.font.name = "Times New Roman"


            szul_ido_text1 = szul_hely_para.add_run("Születési idő:")
            szul_ido_text1.font.bold = True
            szul_ido_text1.font.italic = False
            szul_ido_text1.font.underline = True
            szul_ido_text1.font.size=Pt(12)
            szul_ido_text1.font.name = "Times New Roman"


            szul_ido_text2=szul_hely_para.add_run("\t")
            szul_ido_text2.font.bold=True
            szul_ido_text2.font.italic=False
            szul_ido_text2.font.underline=False
            szul_ido_text2.font.size=Pt(12)
            szul_ido_text2.font.name = "Times New Roman"


            for_year=self.bdate1_V.get()+self.bdate2_V.get()+self.bdate3_V.get()
            try:
                regex=re.match("\d\d\d\d\d\d\d\d", for_year).group()
                szul_ido_text3 = szul_hely_para.add_run(self.bdate1_V.get()+"."+self.bdate2_V.get()+"."+self.bdate3_V.get()+".")
            except:
                szul_ido_text3 = szul_hely_para.add_run("________.____.____.")

            szul_ido_text3.font.bold=True
            szul_ido_text3.font.italic=True
            szul_ido_text3.font.underline=False
            szul_ido_text3.font.size=Pt(12)
            szul_ido_text3.font.name = "Times New Roman"


            szul_hely_para.paragraph_format.space_before = Pt(0)
            szul_hely_para.paragraph_format.space_after = Pt(0)
            szul_hely_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            szul_hely_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            szul_hely_para.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            szul_hely_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)



            allampolgarsag_para=self.save_doc.add_paragraph()
            allampolgarsag_text1=allampolgarsag_para.add_run("Állampolgárság:")
            allampolgarsag_text1.font.bold = True
            allampolgarsag_text1.font.italic = False
            allampolgarsag_text1.font.underline = True
            allampolgarsag_text1.font.size=Pt(12)
            allampolgarsag_text1.font.name = "Times New Roman"

            allampolgarsag_text2=allampolgarsag_para.add_run("\t")
            allampolgarsag_text2.font.bold = True
            allampolgarsag_text2.font.italic = False
            allampolgarsag_text2.font.underline = False
            allampolgarsag_text2.font.size=Pt(12)
            allampolgarsag_text2.font.name = "Times New Roman"


            if str(self.nation_var.get())=="1":
                allampolgarsag_text3 = allampolgarsag_para.add_run("Magyar")
            else:
                allampolgarsag_text3 = allampolgarsag_para.add_run(self.nationality_var.get())
            allampolgarsag_text3.font.bold = True
            allampolgarsag_text3.font.italic = True
            allampolgarsag_text3.font.underline = False
            allampolgarsag_text3.font.size = Pt(12)
            allampolgarsag_text3.font.name = "Times New Roman"


            allampolgarsag_text4=allampolgarsag_para.add_run("\t")
            allampolgarsag_text4.font.bold = True
            allampolgarsag_text4.font.italic = False
            allampolgarsag_text4.font.underline = False
            allampolgarsag_text4.font.size=Pt(12)
            allampolgarsag_text4.font.name = "Times New Roman"


            idegenrend_text1=allampolgarsag_para.add_run("Idegenrend. státusz:")
            idegenrend_text1.font.bold = True
            idegenrend_text1.font.italic = False
            idegenrend_text1.font.underline = True
            idegenrend_text1.font.size=Pt(12)
            idegenrend_text1.font.name = "Times New Roman"


            idegenrend_text2=allampolgarsag_para.add_run("\t")
            idegenrend_text2.font.bold = True
            idegenrend_text2.font.italic = False
            idegenrend_text2.font.underline = False
            idegenrend_text2.font.size=Pt(12)
            idegenrend_text2.font.name = "Times New Roman"


            idegenrend_text3=allampolgarsag_para.add_run(self.idegens[0])
            idegenrend_text3.font.bold = True
            idegenrend_text3.font.italic = True
            idegenrend_text3.font.underline = False
            idegenrend_text3.font.size = Pt(12)
            idegenrend_text3.font.name = "Times New Roman"


            allampolgarsag_para.paragraph_format.space_before = Pt(0)
            allampolgarsag_para.paragraph_format.space_after = Pt(0)
            allampolgarsag_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            allampolgarsag_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            allampolgarsag_para.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            allampolgarsag_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)



            cselekvo_para=self.save_doc.add_paragraph()
            cselekvo_text1=cselekvo_para.add_run("Cselekvőképesség:")
            cselekvo_text1.font.bold = True
            cselekvo_text1.font.italic = False
            cselekvo_text1.font.underline = True
            cselekvo_text1.font.size = Pt(12)
            cselekvo_text1.font.name = "Times New Roman"


            cselekvo_text2=cselekvo_para.add_run("\t")
            cselekvo_text2.font.bold = True
            cselekvo_text2.font.italic = False
            cselekvo_text2.font.underline = False
            cselekvo_text2.font.size=Pt(12)
            cselekvo_text2.font.name = "Times New Roman"

            cselekvo_text3=cselekvo_para.add_run("Cselekvőképes")
            cselekvo_text3.font.bold = True
            cselekvo_text3.font.italic = False
            cselekvo_text3.font.underline = False
            cselekvo_text3.font.size=Pt(12)
            cselekvo_text3.font.name = "Times New Roman"

            cselekvo_text4=cselekvo_para.add_run("\t")
            cselekvo_text4.font.bold = True
            cselekvo_text4.font.italic = False
            cselekvo_text4.font.underline = False
            cselekvo_text4.font.size=Pt(12)
            cselekvo_text4.font.name = "Times New Roman"

            cselekvo_text5 = cselekvo_para.add_run(self.idegens[1])
            cselekvo_text5.font.bold = True
            cselekvo_text5.font.italic = True
            cselekvo_text5.font.underline = False
            cselekvo_text5.font.size = Pt(12)
            cselekvo_text5.font.name = "Times New Roman"

            cselekvo_text6=cselekvo_para.add_run("\t")
            cselekvo_text6.font.bold = True
            cselekvo_text6.font.italic = False
            cselekvo_text6.font.underline = False
            cselekvo_text6.font.size=Pt(12)
            cselekvo_text6.font.name = "Times New Roman"

            cselekvo_text7 = cselekvo_para.add_run(self.idegens[2])
            cselekvo_text7.font.bold = True
            cselekvo_text7.font.italic = True
            cselekvo_text7.font.underline = False
            cselekvo_text7.font.size = Pt(12)
            cselekvo_text7.font.name = "Times New Roman"

            cselekvo_para.paragraph_format.space_before = Pt(5)
            cselekvo_para.paragraph_format.space_after = Pt(10)
            cselekvo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cselekvo_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            cselekvo_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            cselekvo_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)



            folytatas1_para=self.save_doc.add_paragraph()
            folytatas1_text1=folytatas1_para.add_run("\t")
            folytatas1_text1.font.bold = True
            folytatas1_text1.font.italic = False
            folytatas1_text1.font.underline = False
            folytatas1_text1.font.size=Pt(12)
            folytatas1_text1.font.name = "Times New Roman"

            folytatas1_text2=folytatas1_para.add_run("Korlátozottan cselekvőképes")
            folytatas1_text2.font.bold = True
            folytatas1_text2.font.italic = True
            folytatas1_text2.font.underline = False
            folytatas1_text2.font.size=Pt(12)
            folytatas1_text2.font.name = "Times New Roman"
            folytatas1_text3=folytatas1_para.add_run("\t")
            folytatas1_text3.font.bold = True
            folytatas1_text3.font.italic = False
            folytatas1_text3.font.underline = False
            folytatas1_text3.font.size=Pt(12)
            folytatas1_text3.font.name = "Times New Roman"
            folytatas1_text4=folytatas1_para.add_run(self.idegens[3])
            folytatas1_text4.font.bold = True
            folytatas1_text4.font.italic = True
            folytatas1_text4.font.underline = False
            folytatas1_text4.font.size=Pt(12)
            folytatas1_text4.font.name = "Times New Roman"
            folytatas1_text5=folytatas1_para.add_run("\t")
            folytatas1_text5.font.bold = True
            folytatas1_text5.font.italic = False
            folytatas1_text5.font.underline = False
            folytatas1_text5.font.size=Pt(12)
            folytatas1_text5.font.name = "Times New Roman"
            folytatas1_text6=folytatas1_para.add_run(self.idegens[4])
            folytatas1_text6.font.bold = True
            folytatas1_text6.font.italic = True
            folytatas1_text6.font.underline = False
            folytatas1_text6.font.size=Pt(12)
            folytatas1_text6.font.name = "Times New Roman"


            folytatas1_para.paragraph_format.space_before = Pt(5)
            folytatas1_para.paragraph_format.space_after = Pt(0)
            folytatas1_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            folytatas1_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            folytatas1_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            folytatas1_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)


            folytatas2_para=self.save_doc.add_paragraph()
            folytatas2_text1=folytatas2_para.add_run("\t")
            folytatas2_text1.font.bold = True
            folytatas2_text1.font.italic = False
            folytatas2_text1.font.underline = False
            folytatas2_text1.font.size=Pt(12)
            folytatas2_text1.font.name = "Times New Roman"

            folytatas2_text2 = folytatas2_para.add_run("Cselekvőképtelen")
            folytatas2_text2.font.bold = True
            folytatas2_text2.font.italic = True
            folytatas2_text2.font.underline = False
            folytatas2_text2.font.size = Pt(12)
            folytatas2_text2.font.name = "Times New Roman"
            folytatas2_text3 = folytatas2_para.add_run("\t")
            folytatas2_text3.font.bold = True
            folytatas2_text3.font.italic = False
            folytatas2_text3.font.underline = False
            folytatas2_text3.font.size = Pt(12)
            folytatas2_text3.font.name = "Times New Roman"
            folytatas2_text4 = folytatas2_para.add_run("Telefonszám:")
            folytatas2_text4.font.bold = True
            folytatas2_text4.font.italic = False
            folytatas2_text4.font.underline = True
            folytatas2_text4.font.size = Pt(12)
            folytatas2_text4.font.name = "Times New Roman"
            folytatas2_text5 = folytatas2_para.add_run("\t")
            folytatas2_text5.font.bold = True
            folytatas2_text5.font.italic = False
            folytatas2_text5.font.underline = False
            folytatas2_text5.font.size = Pt(12)
            folytatas2_text5.font.name = "Times New Roman"
            for_phone=self.phone_number1.get()+"/ "[0]+self.phone_number2.get()
            try:
                regex=re.match('\d\d\/\d\d\d\d\d\d+', for_phone).group()
                folytatas2_text6 = folytatas2_para.add_run("+36-"+self.phone_number1.get()+"/"+self.phone_number2.get()[:3]+"-"+self.phone_number2.get()[3:])
            except:
                folytatas2_text6 = folytatas2_para.add_run("____\______-_________")

            folytatas2_text6.font.bold = True
            folytatas2_text6.font.italic = True
            folytatas2_text6.font.underline = False
            folytatas2_text6.font.size = Pt(12)
            folytatas2_text6.font.name = "Times New Roman"

            if self.cselekves.get()==1:
                cselekvo_text3.font.underline=True
            elif self.cselekves.get()==2:
                folytatas1_text2.font.underline=True
            else:
                folytatas1_text3.font.underline=True

            if self.idegen.get()==self.idegens[0]:
                idegenrend_text3.font.underline=True
            elif self.idegen.get()==self.idegens[1]:
                cselekvo_text5.font.underline=True
            elif self.idegen.get()==self.idegens[2]:
                cselekvo_text7.font.underline=True
            elif self.idegen.get()==self.idegens[3]:
                folytatas1_text4.font.underline=True
            elif self.idegen.get()==self.idegens[4]:
                folytatas1_text6.font.underline=True



            folytatas2_para.paragraph_format.space_before = Pt(5)
            folytatas2_para.paragraph_format.space_after = Pt(0)
            folytatas2_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            folytatas2_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            folytatas2_para.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            folytatas2_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

            gondnok_para=self.save_doc.add_paragraph()
            gondnok_text1=gondnok_para.add_run("Gondnokolt:")
            gondnok_text1.font.bold = True
            gondnok_text1.font.italic = False
            gondnok_text1.font.underline = True
            gondnok_text1.font.size = Pt(12)
            gondnok_text1.font.name = "Times New Roman"
            gondnok_text2 = gondnok_para.add_run("\t")
            gondnok_text2.font.bold = True
            gondnok_text2.font.italic = False
            gondnok_text2.font.underline = False
            gondnok_text2.font.size = Pt(12)
            gondnok_text2.font.name = "Times New Roman"
            gondnok_text3 = gondnok_para.add_run("Igen")
            gondnok_text3.font.bold = True
            gondnok_text3.font.italic = True
            gondnok_text3.font.underline = False
            gondnok_text3.font.size = Pt(12)
            gondnok_text3.font.name = "Times New Roman"
            gondnok_text4 = gondnok_para.add_run("\t")
            gondnok_text4.font.bold = True
            gondnok_text4.font.italic = False
            gondnok_text4.font.underline = False
            gondnok_text4.font.size = Pt(12)
            gondnok_text4.font.name = "Times New Roman"
            gondnok_text5 = gondnok_para.add_run("Nem")
            gondnok_text5.font.bold = True
            gondnok_text5.font.italic = True
            gondnok_text5.font.underline = False
            gondnok_text5.font.size = Pt(12)
            gondnok_text5.font.name = "Times New Roman"


            if self.gondnok.get() == 1:
                gondnok_text3.font.underline = True
            else:
                gondnok_text5.font.underline = True

            gondnok_para.paragraph_format.tab_stops.add_tab_stop(Inches(1.85), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.SPACES)
            gondnok_para.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.SPACES)
            gondnok_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.25), WD_TAB_ALIGNMENT.LEFT,
                                                                    WD_TAB_LEADER.SPACES)

            gondnok_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.SPACES)



            gondnok_para.paragraph_format.space_before = Pt(5)
            gondnok_para.paragraph_format.space_after = Pt(0)

            gondnoknev_para=self.save_doc.add_paragraph()
            gondnoknev_text1=gondnoknev_para.add_run("Gondnok neve:")
            gondnoknev_text1.font.bold=True
            gondnoknev_text1.font.italic=False
            gondnoknev_text1.font.underline=True
            gondnoknev_text1.font.size = Pt(12)
            gondnoknev_text1.font.name = "Times New Roman"
            gondnoknev_text2=gondnoknev_para.add_run("\t")
            gondnoknev_text2.font.bold=True
            gondnoknev_text2.font.italic=False
            gondnoknev_text2.font.underline=False
            gondnoknev_text2.font.size = Pt(12)
            gondnoknev_text2.font.name = "Times New Roman"
            gondnoknev_text3=gondnoknev_para.add_run(self.gondnok_nev.get())
            gondnoknev_text3.font.bold=True
            gondnoknev_text3.font.italic=True
            gondnoknev_text3.font.underline=True
            gondnoknev_text3.font.size = Pt(12)
            gondnoknev_text3.font.name = "Times New Roman"
            gondnoknev_para.paragraph_format.space_before = Pt(0)
            gondnoknev_para.paragraph_format.space_after = Pt(0)
            gondnoknev_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)

            gondnokcim_para = self.save_doc.add_paragraph()
            gondnokcim_text1 = gondnokcim_para.add_run("Elérhetősége:")
            gondnokcim_text1.font.bold = True
            gondnokcim_text1.font.italic = False
            gondnokcim_text1.font.underline = True
            gondnokcim_text1.font.size = Pt(12)
            gondnokcim_text1.font.name = "Times New Roman"
            gondnokcim_text2 = gondnokcim_para.add_run("\t")
            gondnokcim_text2.font.bold = True
            gondnokcim_text2.font.italic = False
            gondnokcim_text2.font.underline = False
            gondnokcim_text2.font.size = Pt(12)
            gondnokcim_text2.font.name = "Times New Roman"
            gondnokcim_text3 = gondnokcim_para.add_run(self.gondnok_cim.get())
            gondnokcim_text3.font.bold = True
            gondnokcim_text3.font.italic = True
            gondnokcim_text3.font.underline = True
            gondnokcim_text3.font.size = Pt(12)
            gondnokcim_text3.font.name = "Times New Roman"
            gondnokcim_para.paragraph_format.space_before = Pt(0)
            gondnokcim_para.paragraph_format.space_after = Pt(10)
            gondnokcim_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)

            cim_allando_para=self.save_doc.add_paragraph()
            cim_allando_text1=cim_allando_para.add_run("Állandó lakcím:")
            cim_allando_text1.font.bold = True
            cim_allando_text1.font.italic = False
            cim_allando_text1.font.underline = True
            cim_allando_text1.font.size = Pt(14)
            cim_allando_text1.font.name = "Times New Roman"
            cim_allando_para.space_before = Pt(15)
            cim_allando_para.space_after = Pt(0)

            cim_allando_para2=self.save_doc.add_paragraph()
            cim_allando_text3=cim_allando_para2.add_run("Irányítószám:")
            cim_allando_text3.font.bold = True
            cim_allando_text3.font.italic = False
            cim_allando_text3.font.underline = True
            cim_allando_text3.font.size = Pt(12)
            cim_allando_text3.font.name = "Times New Roman"
            cim_allando_text4 = cim_allando_para2.add_run("\t")
            cim_allando_text4.font.bold = True
            cim_allando_text4.font.italic = False
            cim_allando_text4.font.underline = False
            cim_allando_text4.font.size = Pt(12)
            cim_allando_text4.font.name = "Times New Roman"
            cim_allando_text5 = cim_allando_para2.add_run(self.IRSZ_1.get())
            cim_allando_text5.font.bold = True
            cim_allando_text5.font.italic = False
            cim_allando_text5.font.underline = False
            cim_allando_text5.font.size = Pt(12)
            cim_allando_text5.font.name = "Times New Roman"
            cim_allando_text5_5 = cim_allando_para2.add_run("\t")
            cim_allando_text5_5.font.bold = True
            cim_allando_text5_5.font.italic = False
            cim_allando_text5_5.font.underline = False
            cim_allando_text5_5.font.size = Pt(12)
            cim_allando_text5_5.font.name = "Times New Roman"
            cim_allando_text6 = cim_allando_para2.add_run("Város:")
            cim_allando_text6.font.bold = True
            cim_allando_text6.font.italic = False
            cim_allando_text6.font.underline = True
            cim_allando_text6.font.size = Pt(12)
            cim_allando_text7 = cim_allando_para2.add_run("\t")
            cim_allando_text7.font.bold = True
            cim_allando_text7.font.italic = False
            cim_allando_text7.font.underline = False
            cim_allando_text7.font.size = Pt(12)
            cim_allando_text7.font.name = "Times New Roman"
            cim_allando_text8 = cim_allando_para2.add_run(self.city_1.get().replace("(","").replace(",)", "").replace("'",""))
            cim_allando_text8.font.bold = True
            cim_allando_text8.font.italic = False
            cim_allando_text8.font.underline = False
            cim_allando_text8.font.size = Pt(12)
            cim_allando_text8.font.name = "Times New Roman"
            cim_allando_para2.paragraph_format.space_before = Pt(0)
            cim_allando_para2.paragraph_format.space_after = Pt(0)

            cim_allando_para2.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)
            cim_allando_para2.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT,
                                                                    WD_TAB_LEADER.SPACES)
            cim_allando_para2.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)




            cim_allando_para3=self.save_doc.add_paragraph()
            cim_allando_text9=cim_allando_para3.add_run("Cím:")
            cim_allando_text9.font.bold = True
            cim_allando_text9.font.italic = False
            cim_allando_text9.font.underline = True
            cim_allando_text9.font.size = Pt(12)
            cim_allando_text9.font.name = "Times New Roman"
            cim_allando_text10 = cim_allando_para3.add_run("\t")
            cim_allando_text10.font.bold = True
            cim_allando_text10.font.italic = False
            cim_allando_text10.font.underline = False
            cim_allando_text10.font.size = Pt(12)
            cim_allando_text10.font.name = "Times New Roman"

            for_cim=self.ucca1_V.get()
            if self.kozterulet_tipusa.get()=="egyéb:":
                for_cim=for_cim+" "+self.kozterulet_tipusa_egyeb.get()
            else:
                for_cim=for_cim+" "+self.kozterulet_tipusa.get()

            for_cim = for_cim + " " + self.address_V.get()

            cim_allando_text11 = cim_allando_para3.add_run(for_cim)


            cim_allando_text11.font.bold = True
            cim_allando_text11.font.italic = True
            cim_allando_text11.font.underline = False
            cim_allando_text11.font.size = Pt(12)
            cim_allando_text11.font.name = "Times New Roman"

            cim_allando_para2.paragraph_format.space_before = Pt(0)
            cim_allando_para2.paragraph_format.space_after = Pt(15)
            cim_allando_para3.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)

            cim_allando_para.paragraph_format.space_before = Pt(0)
            cim_allando_para.paragraph_format.space_after = Pt(0)
            cim_allando_para2.paragraph_format.space_before = Pt(0)
            cim_allando_para2.paragraph_format.space_after = Pt(0)
            cim_allando_para3.paragraph_format.space_before = Pt(0)
            cim_allando_para3.paragraph_format.space_after = Pt(5)


            cim_tart_para=self.save_doc.add_paragraph()
            cim_tart_text1=cim_tart_para.add_run("Tartózkodási hely:")
            cim_tart_text1.font.bold = True
            cim_tart_text1.font.italic = False
            cim_tart_text1.font.underline = True
            cim_tart_text1.font.size = Pt(14)
            cim_tart_text1.font.name = "Times New Roman"
            cim_tart_para.space_before = Pt(15)
            cim_tart_para.space_after = Pt(0)


            cim_tart_para2=self.save_doc.add_paragraph()
            cim_tart_text3=cim_tart_para2.add_run("Irányítószám:")
            cim_tart_text3.font.bold = True
            cim_tart_text3.font.italic = False
            cim_tart_text3.font.underline = True
            cim_tart_text3.font.size = Pt(12)
            cim_tart_text3.font.name = "Times New Roman"
            cim_tart_text4 = cim_tart_para2.add_run("\t")
            cim_tart_text4.font.bold = True
            cim_tart_text4.font.italic = False
            cim_tart_text4.font.underline = False
            cim_tart_text4.font.size = Pt(12)
            cim_tart_text4.font.name = "Times New Roman"
            cim_tart_text5 = cim_tart_para2.add_run(self.IRSZ_2.get())
            cim_tart_text5.font.bold = True
            cim_tart_text5.font.italic = False
            cim_tart_text5.font.underline = False
            cim_tart_text5.font.size = Pt(12)
            cim_tart_text5.font.name = "Times New Roman"
            cim_tart_text5_5 = cim_tart_para2.add_run("\t")
            cim_tart_text5_5.font.bold = True
            cim_tart_text5_5.font.italic = False
            cim_tart_text5_5.font.underline = False
            cim_tart_text5_5.font.size = Pt(12)
            cim_tart_text5_5.font.name = "Times New Roman"
            cim_tart_text6 = cim_tart_para2.add_run("Város:")
            cim_tart_text6.font.bold = True
            cim_tart_text6.font.italic = False
            cim_tart_text6.font.underline = True
            cim_tart_text6.font.size = Pt(12)
            cim_tart_text7 = cim_tart_para2.add_run("\t")
            cim_tart_text7.font.bold = True
            cim_tart_text7.font.italic = False
            cim_tart_text7.font.underline = False
            cim_tart_text7.font.size = Pt(12)
            cim_tart_text7.font.name = "Times New Roman"
            cim_tart_text8 = cim_tart_para2.add_run(self.city_2.get().replace("(","").replace(",)", "").replace("'",""))
            cim_tart_text8.font.bold = True
            cim_tart_text8.font.italic = False
            cim_tart_text8.font.underline = False
            cim_tart_text8.font.size = Pt(12)
            cim_tart_text8.font.name = "Times New Roman"
            cim_tart_para2.paragraph_format.space_before = Pt(0)
            cim_tart_para2.paragraph_format.space_after = Pt(0)

            cim_tart_para2.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)
            cim_tart_para2.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT,
                                                                    WD_TAB_LEADER.SPACES)
            cim_tart_para2.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)



            cim_tart_para3=self.save_doc.add_paragraph()
            cim_tart_text9=cim_tart_para3.add_run("Cím:")
            cim_tart_text9.font.bold = True
            cim_tart_text9.font.italic = False
            cim_tart_text9.font.underline = True
            cim_tart_text9.font.size = Pt(12)
            cim_tart_text9.font.name = "Times New Roman"
            cim_tart_text10 = cim_tart_para3.add_run("\t")
            cim_tart_text10.font.bold = True
            cim_tart_text10.font.italic = False
            cim_tart_text10.font.underline = False
            cim_tart_text10.font.size = Pt(12)
            cim_tart_text10.font.name = "Times New Roman"

            for_cim=self.ucca2_V.get()
            if self.kozterulet_tipusa2.get()=="egyéb:":
                for_cim=for_cim+" "+self.kozterulet_tipusa2_egyeb.get()
            else:
                for_cim=for_cim+" "+self.kozterulet_tipusa2.get()

            for_cim = for_cim + " " + self.address2_V.get()

            cim_tart_text11 = cim_tart_para3.add_run(for_cim)


            cim_tart_text11.font.bold = True
            cim_tart_text11.font.italic = True
            cim_tart_text11.font.underline = False
            cim_tart_text11.font.size = Pt(12)
            cim_tart_text11.font.name = "Times New Roman"

            cim_tart_para2.paragraph_format.space_before = Pt(0)
            cim_tart_para2.paragraph_format.space_after = Pt(15)
            cim_tart_para3.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)

            cim_tart_para.paragraph_format.space_before = Pt(0)
            cim_tart_para.paragraph_format.space_after = Pt(0)
            cim_tart_para2.paragraph_format.space_before = Pt(0)
            cim_tart_para2.paragraph_format.space_after = Pt(0)
            cim_tart_para3.paragraph_format.space_before = Pt(0)
            cim_tart_para3.paragraph_format.space_after = Pt(5)


            cim_eletvitel_para=self.save_doc.add_paragraph()
            cim_eletvitel_text1=cim_eletvitel_para.add_run("Életvitelszerű tartózkodási hely:")
            cim_eletvitel_text1.font.bold = True
            cim_eletvitel_text1.font.italic = False
            cim_eletvitel_text1.font.underline = True
            cim_eletvitel_text1.font.size = Pt(14)
            cim_eletvitel_text1.font.name = "Times New Roman"
            cim_eletvitel_para.space_before = Pt(15)
            cim_eletvitel_para.space_after = Pt(0)


            cim_eletvitel_para2=self.save_doc.add_paragraph()
            cim_eletvitel_text3=cim_eletvitel_para2.add_run("Irányítószám:")
            cim_eletvitel_text3.font.bold = True
            cim_eletvitel_text3.font.italic = False
            cim_eletvitel_text3.font.underline = True
            cim_eletvitel_text3.font.size = Pt(12)
            cim_eletvitel_text3.font.name = "Times New Roman"
            cim_eletvitel_text4 = cim_eletvitel_para2.add_run("\t")
            cim_eletvitel_text4.font.bold = True
            cim_eletvitel_text4.font.italic = False
            cim_eletvitel_text4.font.underline = False
            cim_eletvitel_text4.font.size = Pt(12)
            cim_eletvitel_text4.font.name = "Times New Roman"
            cim_eletvitel_text5 = cim_eletvitel_para2.add_run(self.IRSZ_3.get())
            cim_eletvitel_text5.font.bold = True
            cim_eletvitel_text5.font.italic = False
            cim_eletvitel_text5.font.underline = False
            cim_eletvitel_text5.font.size = Pt(12)
            cim_eletvitel_text5.font.name = "Times New Roman"
            cim_eletvitel_text5_5 = cim_eletvitel_para2.add_run("\t")
            cim_eletvitel_text5_5.font.bold = True
            cim_eletvitel_text5_5.font.italic = False
            cim_eletvitel_text5_5.font.underline = False
            cim_eletvitel_text5_5.font.size = Pt(12)
            cim_eletvitel_text5_5.font.name = "Times New Roman"
            cim_eletvitel_text6 = cim_eletvitel_para2.add_run("Város:")
            cim_eletvitel_text6.font.bold = True
            cim_eletvitel_text6.font.italic = False
            cim_eletvitel_text6.font.underline = True
            cim_eletvitel_text6.font.size = Pt(12)
            cim_eletvitel_text7 = cim_eletvitel_para2.add_run("\t")
            cim_eletvitel_text7.font.bold = True
            cim_eletvitel_text7.font.italic = False
            cim_eletvitel_text7.font.underline = False
            cim_eletvitel_text7.font.size = Pt(12)
            cim_eletvitel_text7.font.name = "Times New Roman"
            cim_eletvitel_text8 = cim_eletvitel_para2.add_run(self.city_3.get().replace("(","").replace(",)", "").replace("'",""))
            cim_eletvitel_text8.font.bold = True
            cim_eletvitel_text8.font.italic = False
            cim_eletvitel_text8.font.underline = False
            cim_eletvitel_text8.font.size = Pt(12)
            cim_eletvitel_text8.font.name = "Times New Roman"
            cim_eletvitel_para2.paragraph_format.space_before = Pt(0)
            cim_eletvitel_para2.paragraph_format.space_after = Pt(0)

            cim_eletvitel_para2.paragraph_format.tab_stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)
            cim_eletvitel_para2.paragraph_format.tab_stops.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.LEFT,
                                                                    WD_TAB_LEADER.SPACES)
            cim_eletvitel_para2.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)



            cim_eletvitel_para3=self.save_doc.add_paragraph()
            cim_eletvitel_text9=cim_eletvitel_para3.add_run("Cím:")
            cim_eletvitel_text9.font.bold = True
            cim_eletvitel_text9.font.italic = False
            cim_eletvitel_text9.font.underline = True
            cim_eletvitel_text9.font.size = Pt(12)
            cim_eletvitel_text9.font.name = "Times New Roman"
            cim_eletvitel_text10 = cim_eletvitel_para3.add_run("\t")
            cim_eletvitel_text10.font.bold = True
            cim_eletvitel_text10.font.italic = False
            cim_eletvitel_text10.font.underline = False
            cim_eletvitel_text10.font.size = Pt(12)
            cim_eletvitel_text10.font.name = "Times New Roman"

            for_cim=self.ucca2_V.get()
            if self.kozterulet_tipusa2.get()=="egyéb:":
                for_cim=for_cim+" "+self.kozterulet_tipusa3_egyeb.get()
            else:
                for_cim=for_cim+" "+self.kozterulet_tipusa3.get()

            for_cim = for_cim + " " + self.address3_V.get()

            cim_eletvitel_text11 = cim_eletvitel_para3.add_run(for_cim)


            cim_eletvitel_text11.font.bold = True
            cim_eletvitel_text11.font.italic = True
            cim_eletvitel_text11.font.underline = False
            cim_eletvitel_text11.font.size = Pt(12)
            cim_eletvitel_text11.font.name = "Times New Roman"

            cim_eletvitel_para2.space_before = Pt(0)
            cim_eletvitel_para2.space_after = Pt(15)
            cim_eletvitel_para3.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                    WD_TAB_LEADER.DOTS)


            cim_eletvitel_para.paragraph_format.space_before = Pt(0)
            cim_eletvitel_para.paragraph_format.space_after = Pt(0)
            cim_eletvitel_para2.paragraph_format.space_before = Pt(0)
            cim_eletvitel_para2.paragraph_format.space_after = Pt(0)
            cim_eletvitel_para3.paragraph_format.space_before = Pt(0)
            cim_eletvitel_para3.paragraph_format.space_after = Pt(5)

            kozeli_para = self.save_doc.add_paragraph()
            kozeli_text1 = kozeli_para.add_run("Megnevezett nagykorú hozzátartozó:")
            kozeli_text1.font.bold = True
            kozeli_text1.font.italic = False
            kozeli_text1.font.underline = True
            kozeli_text1.font.size = Pt(12)
            kozeli_text1.font.name = "Times New Roman"
            kozeli_text2 = kozeli_para.add_run("\t")
            kozeli_text2.font.bold = True
            kozeli_text2.font.italic = False
            kozeli_text2.font.underline = False
            kozeli_text2.font.size = Pt(12)
            kozeli_text2.font.name = "Times New Roman"
            kozeli_text2 = kozeli_para.add_run(self.hozzatart_V.get())
            kozeli_text2.font.bold = True
            kozeli_text2.font.italic = True
            kozeli_text2.font.underline = False
            kozeli_text2.font.size = Pt(12)
            kozeli_text2.font.name = "Times New Roman"
            kozeli_para.paragraph_format.space_before = Pt(0)
            kozeli_para.paragraph_format.space_after = Pt(0)

            kozeli_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

            kozeli_eler_para = self.save_doc.add_paragraph()
            kozeli_eler_text1 = kozeli_eler_para.add_run("Hozzátartozó elérhetősége:")
            kozeli_eler_text1.font.bold = True
            kozeli_eler_text1.font.italic = False
            kozeli_eler_text1.font.underline = True
            kozeli_eler_text1.font.size = Pt(12)
            kozeli_eler_text1.font.name = "Times New Roman"
            kozeli_eler_text2 = kozeli_eler_para.add_run("\t")
            kozeli_eler_text2.font.bold = True
            kozeli_eler_text2.font.italic = False
            kozeli_eler_text2.font.underline = False
            kozeli_eler_text2.font.size = Pt(12)
            kozeli_eler_text2.font.name = "Times New Roman"
            kozeli_eler_text2 = kozeli_eler_para.add_run(self.eler_V.get())
            kozeli_eler_text2.font.bold = True
            kozeli_eler_text2.font.italic = True
            kozeli_eler_text2.font.underline = False
            kozeli_eler_text2.font.size = Pt(12)
            kozeli_eler_text2.font.name = "Times New Roman"
            kozeli_eler_para.paragraph_format.space_before = Pt(0)
            kozeli_eler_para.paragraph_format.space_after = Pt(0)

            kozeli_eler_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT,
                                                                     WD_TAB_LEADER.DOTS)

            soron_kivul_para=self.save_doc.add_paragraph()
            soron_kivul_text1=soron_kivul_para.add_run("Soron kívül elhelyezés iránti kérelem:")
            soron_kivul_text1.font.bold=True
            soron_kivul_text1.font.italic=False
            soron_kivul_text1.font.underline = True
            soron_kivul_text1.font.size = Pt(12)
            soron_kivul_text1.font.name = "Times New Roman"


            soron_kivul_text2=soron_kivul_para.add_run("\t")
            soron_kivul_text2.font.bold=True
            soron_kivul_text2.font.italic=False
            soron_kivul_text2.font.underline = False
            soron_kivul_text2.font.size = Pt(12)
            soron_kivul_text2.font.name = "Times New Roman"

            soron_kivul_text3=soron_kivul_para.add_run("Igen")
            soron_kivul_text3.font.bold=True
            soron_kivul_text3.font.italic=True
            soron_kivul_text3.font.underline = False
            soron_kivul_text3.font.size = Pt(12)
            soron_kivul_text3.font.name = "Times New Roman"

            soron_kivul_text4=soron_kivul_para.add_run("\t")
            soron_kivul_text4.font.bold=True
            soron_kivul_text4.font.italic=False
            soron_kivul_text4.font.underline = False
            soron_kivul_text4.font.size = Pt(12)
            soron_kivul_text4.font.name = "Times New Roman"

            soron_kivul_text5=soron_kivul_para.add_run("Nem")
            soron_kivul_text5.font.bold=True
            soron_kivul_text5.font.italic=True
            soron_kivul_text5.font.underline = False
            soron_kivul_text5.font.size = Pt(12)
            soron_kivul_text5.font.name = "Times New Roman"

            soron_kivul_para.paragraph_format.space_before = Pt(0)
            soron_kivul_para.paragraph_format.space_after = Pt(0)
            soron_kivul_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.RIGHT,
                                                                        WD_TAB_LEADER.SPACES)
            soron_kivul_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT,
                                                                        WD_TAB_LEADER.SPACES)

            if self.soron_kivul_V.get()==1:
                soron_kivul_text3.font.underline = True
            else:
                soron_kivul_text5.font.underline = True

            elogondozas_para = self.save_doc.add_paragraph()
            elogondozas_text1 = elogondozas_para.add_run("Előgondozás történt:")
            elogondozas_text1.font.bold = True
            elogondozas_text1.font.italic = False
            elogondozas_text1.font.underline = True
            elogondozas_text1.font.size = Pt(12)
            elogondozas_text1.font.name = "Times New Roman"

            elogondozas_text2 = elogondozas_para.add_run("\t")
            elogondozas_text2.font.bold = True
            elogondozas_text2.font.italic = False
            elogondozas_text2.font.underline = False
            elogondozas_text2.font.size = Pt(12)
            elogondozas_text2.font.name = "Times New Roman"

            elogondozas_text3 = elogondozas_para.add_run("Igen")
            elogondozas_text3.font.bold = True
            elogondozas_text3.font.italic = True
            elogondozas_text3.font.underline = False
            elogondozas_text3.font.size = Pt(12)
            elogondozas_text3.font.name = "Times New Roman"

            elogondozas_text4 = elogondozas_para.add_run("\t")
            elogondozas_text4.font.bold = True
            elogondozas_text4.font.italic = False
            elogondozas_text4.font.underline = False
            elogondozas_text4.font.size = Pt(12)
            elogondozas_text4.font.name = "Times New Roman"

            elogondozas_text5 = elogondozas_para.add_run("Nem")
            elogondozas_text5.font.bold = True
            elogondozas_text5.font.italic = True
            elogondozas_text5.font.underline = False
            elogondozas_text5.font.size = Pt(12)
            elogondozas_text5.font.name = "Times New Roman"

            elogondozas_para.paragraph_format.space_before = Pt(0)
            elogondozas_para.paragraph_format.space_after = Pt(0)
            elogondozas_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.RIGHT,
                                                                     WD_TAB_LEADER.SPACES)
            elogondozas_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT,
                                                                     WD_TAB_LEADER.SPACES)

            if self.elogondozás_V.get() == 1:
                elogondozas_text3.font.underline = True
            else:
                elogondozas_text5.font.underline = True

            elogondozas_date_para = self.save_doc.add_paragraph()
            elogondozas_date_text1 = elogondozas_date_para.add_run("Előgondozás dátuma:")
            elogondozas_date_text1.font.bold = True
            elogondozas_date_text1.font.italic = False
            elogondozas_date_text1.font.underline = True
            elogondozas_date_text1.font.size = Pt(12)
            elogondozas_date_text1.font.name = "Times New Roman"

            elogondozas_date_text2 = elogondozas_date_para.add_run("\t")
            elogondozas_date_text2.font.bold = True
            elogondozas_date_text2.font.italic = False
            elogondozas_date_text2.font.underline = False
            elogondozas_date_text2.font.size = Pt(12)
            elogondozas_date_text2.font.name = "Times New Roman"

            for_date=self.edate1_V.get()+"."+self.edate2_V.get()+"."+self.edate3_V.get()+"."
            try:
                regex=re.match('\d\d\d\d\.\d\d?\.\d\d?\.', for_date).group()
                elogondozas_date_text3 = elogondozas_date_para.add_run(for_date)

            except:
                elogondozas_date_text3 = elogondozas_date_para.add_run("________.____.____.")
            elogondozas_date_text3.font.bold = True
            elogondozas_date_text3.font.italic = True
            elogondozas_date_text3.font.underline = False
            elogondozas_date_text3.font.size = Pt(12)
            elogondozas_date_text3.font.name = "Times New Roman"

            keltezes_para=self.save_doc.add_paragraph()
            keltezes_text1=keltezes_para.add_run("______________________________,")
            for_date=self.kdate1_V.get()+"."+self.kdate2_V.get()+"."+self.kdate3_V.get()+"."
            try:
                regex=re.match('\d\d\d\d\.\d\d?\.\d\d?\.', for_date).group()
                keltezes_text2 = keltezes_para.add_run(for_date)

            except:
                keltezes_text2 = keltezes_para.add_run("________.____.____.")
            keltezes_text2.font.bold = True
            keltezes_text2.font.italic = True
            keltezes_text2.font.underline = False
            keltezes_text2.font.size = Pt(12)
            keltezes_text2.font.name = "Times New Roman"





            self.save_doc.save(self.filename)

    def search(self):
        self.search_box = Toplevel()
        self.search_box.title("Ügyfél Keresése")
        root_x = Ablak.winfo_x()
        root_y = Ablak.winfo_y()
        root_width = Ablak.winfo_width()
        root_height = Ablak.winfo_height()
        self.search_box.resizable(False, False)
        self.search_box.geometry("%dx%d+%d+%d" % (450, 195, root_x + root_width / 2 - 220, root_y + root_height / 2 - 60))
        self.scroller2 = Scrollbar(self.search_box)
        self.list_of_people2 = Listbox(self.search_box, yscrollcommand=self.scroller2.set, width=30)
        self.scroller2.config(command=self.list_of_people2.yview)
        self.active_name2 = self.list_of_people2.get(ACTIVE)
        self.now2 = self.list_of_people2.curselection()

        self.list_of_people2.pack(side=LEFT, fill=Y, anchor=W)
        self.scroller2.pack(side=LEFT, fill=BOTH, anchor=W)
        self.text29=Label(self.search_box, text="Ügyfél neve:")
        self.text29.place(x=210, y=5)
        self.search_name_V=StringVar()
        self.search_name_E=Entry(self.search_box, width=21, textvariable=self.search_name_V)
        self.search_name_E.place(x=310, y=5)
        self.text30=Label(self.search_box, text="Születési dátum:")
        self.text30.place(x=210, y=30)
        self.search_bdate1_V=StringVar()
        self.search_bdate2_V=StringVar()
        self.search_bdate3_V=StringVar()
        self.search_bdate1_E=Entry(self.search_box, width=4, textvariable=self.search_bdate1_V, justify=RIGHT)
        self.search_bdate2_E=Entry(self.search_box, width=2, textvariable=self.search_bdate2_V, justify=RIGHT)
        self.search_bdate3_E=Entry(self.search_box, width=2, textvariable=self.search_bdate3_V, justify=RIGHT)
        self.search_bline1=Label(self.search_box, text="-")
        self.search_bline2=Label(self.search_box, text="-")
        self.search_bdate1_E.place(x=310, y=30)
        self.search_bline1.place(x=340, y=30)
        self.search_bdate2_E.place(x=355, y=30)
        self.search_bline2.place(x=370, y=30)
        self.search_bdate3_E.place(x=385, y=30)
        self.text31=Label(self.search_box, text="TAJ szám:")
        self.text31.place(x=210, y=55)
        self.search_TAJ1_V=StringVar()
        self.search_TAJ2_V=StringVar()
        self.search_TAJ3_V=StringVar()
        self.search_TAJ4_V=StringVar()

        self.search_TAJ1_E=Entry(self.search_box, width=3, textvariable=self.search_TAJ1_V, justify=RIGHT)
        self.search_TAJ2_E=Entry(self.search_box, width=3, textvariable=self.search_TAJ2_V, justify=RIGHT)
        self.search_TAJ3_E=Entry(self.search_box, width=2, textvariable=self.search_TAJ3_V, justify=RIGHT)
        self.search_TAJ4_E=Entry(self.search_box, width=1, textvariable=self.search_TAJ4_V, state=DISABLED)
        self.search_TAJ_line1=Label(self.search_box, text="-")
        self.search_TAJ_line2=Label(self.search_box, text="-")
        self.search_TAJ1_E.place(x=310, y=55)
        self.search_TAJ_line1.place(x=330, y=55)
        self.search_TAJ2_E.place(x=340, y=55)
        self.search_TAJ_line2.place(x=360, y=55)
        self.search_TAJ3_E.place(x=370, y=55)
        self.search_TAJ4_E.place(x=385, y=55)
        self.text32=Label(self.search_box, text="Anyja neve:")
        self.text32.place(x=210, y=80)
        self.search_mothername_V=StringVar()
        self.search_mothername_E=Entry(self.search_box, width=21, textvariable=self.search_mothername_V)
        self.search_mothername_E.place(x=310, y=80)
        self.LoadIt=Button(self.search_box, width=20, text="Személy betöltés", command=self.load_it)
        self.LoadIt.place(x=290, y=160)
        self.where=IntVar()
        self.wherebutton1=Radiobutton(self.search_box, text="Átmeneti Szállás", variable=self.where, value=1)
        self.wherebutton2=Radiobutton(self.search_box, text="Éjjeli Menedékhely", variable=self.where, value=2)
        self.wherebutton3=Radiobutton(self.search_box, text="Nappali Melegedő", variable=self.where, value=3)
        self.wherebutton1.place(x=205, y=110)
        self.wherebutton2.place(x=310, y=110)
        self.wherebutton3.place(x=205, y=130)


        self.search_box_check()

    def load_it(self):
        try:
            self.active_name=self.active_name2
            self.fill_table()
        except:
            pass

    def search_box_check(self):
        self.active_name2 = self.list_of_people2.get(ACTIVE)
        if self.list_of_people.curselection() != self.now:
            pass

        self.now2 = self.list_of_people2.curselection()

        self.search_bdate()
        self.search_TAJ_check()
        self.search_list()
        self.search_box.after(250, self.search_box_check)

    def search_list(self):
        self.list_of_people2.delete(0, END)
        if self.where.get()==1:
            relative=self.ugyfellista
        elif self.where.get()==2:
            relative=self.ugyfellista_ejjeli
        else:
            relative=self.ugyfellista_nappali


        for ember in relative:
            add=False
            if self.search_name_V.get()!="" and self.search_name_V.get()==relative[ember][0]:
                add=True
            elif self.search_bdate1_V.get()!="" and self.search_bdate2_V.get()!="" and self.search_bdate3_V.get()!="" and self.search_bdate1_V.get()==relative[ember][10] and self.search_bdate2_V.get()==relative[ember][11] and self.search_bdate3_V.get()==relative[ember][12]:
                add=True
            elif self.search_TAJ1_V.get()!="" and self.search_TAJ2_V.get()!="" and self.search_TAJ3_V.get()!="" and self.search_TAJ1_V.get()==relative[ember][1] and self.search_TAJ2_V.get()==relative[ember][2] and self.search_TAJ3_V.get()==relative[ember][3]:
                add=True
            elif self.search_mothername_V.get()!="" and self.search_mothername_V.get()==relative[ember][48]:
                add=True
            if add==True:
                self.list_of_people2.insert(END, str(ember+".: "+relative[ember][0]+" ("+relative[ember][10]+")"))

    def search_TAJ_check(self):
        self.search_TAJ1_E.config(bg="white", fg="black")
        self.search_TAJ2_E.config(bg="white", fg="black")
        self.search_TAJ3_E.config(bg="white", fg="black")
        self.search_TAJ1_V.set(self.search_TAJ1_V.get()[:3])
        self.search_TAJ2_V.set(self.search_TAJ2_V.get()[:3])
        self.search_TAJ3_V.set(self.search_TAJ3_V.get()[:2])

        try:
            teszt=int(self.search_TAJ1_V.get())
        except:
            self.search_TAJ1_E.config(bg="red", fg="white")
        try:
            teszt=int(self.search_TAJ2_V.get())
        except:
            self.search_TAJ2_E.config(bg="red", fg="white")
        try:
            teszt=int(self.search_TAJ3_V.get())
        except:
            self.search_TAJ3_E.config(bg="red", fg="white")
        try:
            self.search_TAJ4_E.configure(state=NORMAL)
            self.search_TAJ4_V.set(str((int(self.search_TAJ1_V.get()[0])*3+int(self.search_TAJ1_V.get()[1])*7+int(self.search_TAJ1_V.get()[2])*3+int(self.search_TAJ2_V.get()[0])*7+int(self.search_TAJ2_V.get()[1])*3+int(self.search_TAJ2_V.get()[2])*7+int(self.search_TAJ3_V.get()[0])*3+int(self.search_TAJ3_V.get()[1])*7)%10))
            if self.search_TAJ4_V.get()=="10":
                self.search_TAJ4_V.set("0")
            self.search_TAJ4_E.configure(state=DISABLED)

        except:
            pass

    def search_bdate(self):
        self.search_bdate1_E.config(bg="white", fg="black")
        self.search_bdate2_E.config(bg="white", fg="black")
        self.search_bdate3_E.config(bg="white", fg="black")
        self.search_bdate1_V.set(self.search_bdate1_V.get()[:4])
        self.search_bdate2_V.set(self.search_bdate2_V.get()[:2])
        self.search_bdate3_V.set(self.search_bdate3_V.get()[:2])

        try:
            teszt=int(self.search_bdate1_V.get())
        except:
            self.search_bdate1_E.config(bg="red", fg="white")
        try:
            teszt=int(self.search_bdate2_V.get())
        except:
            self.search_bdate2_E.config(bg="red", fg="white")
        try:
            teszt=int(self.search_bdate3_V.get())
        except:
            self.search_bdate3_E.config(bg="red", fg="white")

        try:
            if int(self.search_bdate2_V.get())>12:
                self.search_bdate2_E.config(bg="red", fg="white")
        except:
            pass

        try:
            if int(self.search_bdate3_V.get())>31:
                self.search_bdate3_E.config(bg="red", fg="white")
        except:
            pass

        try:
            if self.search_bdate2_V.get() in ["4", "04", "6", "06", "9", "09", "11"] and int(self.bdate3_V.get()>30):
                self.search_bdate3_E.config(bg="red", fg="white")
        except:
            pass

        if self.search_bdate2_V.get() in ["2", "02"]:
            try:
                if int(self.search_bdate1_V.get()) % 4 != 0 and int(self.search_bdate1_V.get()) % 400 != 0 and int(
                        self.search_bdate3_V.get()) > 28:
                    self.search_bdate3_E.config(bg="red", fg="white")

                elif self.search_bdate1_V.get()[2:] == "00" and int(self.search_bdate1_V.get()[:2]) % 4 != 0 and int(
                        self.search_bdate3_V.get()) > 28:
                    self.search_bdate3_E.config(bg="red", fg="white")

                elif int(self.search_bdate3_V.get()) > 29:
                    self.search_bdate3_E.config(bg="red", fg="white")
            except:
                pass


    def save_person(self):
        if self.forma.get() == self.formak[0]:
            relative = self.ugyfellista
        elif self.forma.get() == self.formak[1]:
            relative = self.ugyfellista_ejjeli
        else:
            relative = self.ugyfellista_nappali

        if self.name_var.get=="" or self.bdate1_V.get()=="":
            msg=messagebox.showerror("Hiba!", "Legalább a név és születési év megadása kötelező!")
            return
        if str(self.person_number.get()) in relative:
            msg=messagebox.askquestion("Már létező személy!", "A megadott személy már létezik az adatbázisban! Felülírja?")
            if msg=="no":
                return
        else:
            relative.setdefault(str(self.person_number.get()), [])
            for ures in range(0,52):
                relative[str(self.person_number.get())].append("")
            to_be_inserted = str(
                str(self.person_number.get()) + ".: " + self.name_var.get() + " (" + self.bdate1_V.get() + ")")
            self.list_of_people.insert(END, to_be_inserted)
            self.list_of_people.see(END)
        #if self.error==True:
            #msg = messagebox.askquestion("Hiba az adatlapon!", "A menteni kívánt adatlapon hiba van! Kívánja menteni?")
            #if msg == "no":
                #return

        relative[str(self.person_number.get())][0]=self.name_var.get()
        relative[str(self.person_number.get())][1]=self.TAJ1_V.get()
        relative[str(self.person_number.get())][2]=self.TAJ2_V.get()
        relative[str(self.person_number.get())][3]=self.TAJ3_V.get()
        relative[str(self.person_number.get())][4]=self.TAJ4_V.get()
        relative[str(self.person_number.get())][5] = self.nincsTAJ.get()
        relative[str(self.person_number.get())][6]=self.bname_var.get()
        relative[str(self.person_number.get())][7]=self.same_V.get()
        relative[str(self.person_number.get())][8]=self.nation_var.get()
        relative[str(self.person_number.get())][9]=self.nationality_var.get()
        relative[str(self.person_number.get())][10]=self.bplace_V.get()
        relative[str(self.person_number.get())][11]=self.bdate1_V.get()
        relative[str(self.person_number.get())][12]=self.bdate2_V.get()
        relative[str(self.person_number.get())][13]=self.bdate3_V.get()
        relative[str(self.person_number.get())][14]=self.phone_number1.get()
        relative[str(self.person_number.get())][15]=self.phone_number2.get()
        relative[str(self.person_number.get())][16]=self.idegen.get()
        relative[str(self.person_number.get())][17]=self.cselekves.get()
        relative[str(self.person_number.get())][18]=self.gondnok.get()
        relative[str(self.person_number.get())][19]=self.IRSZ_1.get()
        relative[str(self.person_number.get())][20]=self.city_1.get()
        relative[str(self.person_number.get())][21]=self.ucca1_V.get()
        relative[str(self.person_number.get())][22]=self.kozterulet_tipusa.get()
        relative[str(self.person_number.get())][23]=self.kozterulet_tipusa_egyeb.get()
        relative[str(self.person_number.get())][24]=self.address_V.get()
        relative[str(self.person_number.get())][25]=self.IRSZ_2.get()
        relative[str(self.person_number.get())][26]=self.city_2.get()
        relative[str(self.person_number.get())][27]=self.ucca2_V.get()
        relative[str(self.person_number.get())][28]=self.kozterulet_tipusa2.get()
        relative[str(self.person_number.get())][29]=self.kozterulet_tipusa2_egyeb.get()
        relative[str(self.person_number.get())][30]=self.address2_V.get()
        relative[str(self.person_number.get())][31]=self.IRSZ_3.get()
        relative[str(self.person_number.get())][32]=self.city_3.get()
        relative[str(self.person_number.get())][33]=self.ucca3_V.get()
        relative[str(self.person_number.get())][34]=self.kozterulet_tipusa3.get()
        relative[str(self.person_number.get())][35]=self.kozterulet_tipusa3_egyeb.get()
        relative[str(self.person_number.get())][36]=self.address3_V.get()
        relative[str(self.person_number.get())][37]=self.megegyezik1.get()
        relative[str(self.person_number.get())][38]=self.megegyezik2.get()
        relative[str(self.person_number.get())][39]=self.hozzatart_V.get()
        relative[str(self.person_number.get())][40]=self.eler_V.get()
        relative[str(self.person_number.get())][41]=self.soron_kivul_V.get()
        relative[str(self.person_number.get())][42]=self.elogondozás_V.get()
        relative[str(self.person_number.get())][43]=self.kdate1_V.get()
        relative[str(self.person_number.get())][44]=self.kdate2_V.get()
        relative[str(self.person_number.get())][45]=self.kdate3_V.get()
        relative[str(self.person_number.get())][46]=self.edate1_V.get()
        relative[str(self.person_number.get())][47]=self.edate2_V.get()
        relative[str(self.person_number.get())][48]=self.edate3_V.get()
        relative[str(self.person_number.get())][49]=self.mothername.get()
        relative[str(self.person_number.get())][50]=self.gondnok_nev.get()
        relative[str(self.person_number.get())][51]=self.gondnok_cim.get()

        self.changed=True
        self.zero_record()


    def fill_table(self):
        if  self.forma.get()==self.formak[0]:
            relative=self.ugyfellista
        elif self.forma.get()==self.formak[1]:
            relative = self.ugyfellista_ejjeli
        else:
            relative = self.ugyfellista_nappali
        try:
            self.person_number.set(int((self.active_name).split(".")[0]))
            self.name_var.set(relative[str(self.person_number.get())][0])
            self.TAJ1_V.set(relative[str(self.person_number.get())][1])
            self.TAJ2_V.set(relative[str(self.person_number.get())][2])
            self.TAJ3_V.set(relative[str(self.person_number.get())][3])
            self.TAJ4_V.set(relative[str(self.person_number.get())][4])
            self.nincsTAJ.set(relative[str(self.person_number.get())][5])
            self.bname_var.set(relative[str(self.person_number.get())][6])
            self.same_V.set(relative[str(self.person_number.get())][7])
            self.nation_var.set(relative[str(self.person_number.get())][8])
            self.nationality_var.set(relative[str(self.person_number.get())][9])
            self.bplace_V.set(relative[str(self.person_number.get())][10])
            self.bdate1_V.set(relative[str(self.person_number.get())][11])
            self.bdate2_V.set(relative[str(self.person_number.get())][12])
            self.bdate3_V.set(relative[str(self.person_number.get())][13])
            self.phone_number1.set(relative[str(self.person_number.get())][14])
            self.phone_number2.set(relative[str(self.person_number.get())][15])
            self.idegen.set(relative[str(self.person_number.get())][16])
            self.cselekves.set(relative[str(self.person_number.get())][17])
            self.gondnok.set(relative[str(self.person_number.get())][18])
            self.IRSZ_1.set(relative[str(self.person_number.get())][19])
            self.city_1.set(relative[str(self.person_number.get())][20])
            self.ucca1_V.set(relative[str(self.person_number.get())][21])
            self.kozterulet_tipusa.set(relative[str(self.person_number.get())][22])
            self.kozterulet_tipusa_egyeb.set(relative[str(self.person_number.get())][23])
            self.address_V.set(relative[str(self.person_number.get())][24])
            self.IRSZ_2.set(relative[str(self.person_number.get())][25])
            self.city_2.set(relative[str(self.person_number.get())][26])
            self.ucca2_V.set(relative[str(self.person_number.get())][27])
            self.kozterulet_tipusa2.set(relative[str(self.person_number.get())][28])
            self.kozterulet_tipusa2_egyeb.set(relative[str(self.person_number.get())][29])
            self.address2_V.set(relative[str(self.person_number.get())][30])
            self.IRSZ_3.set(relative[str(self.person_number.get())][31])
            self.city_3.set(relative[str(self.person_number.get())][32])
            self.ucca3_V.set(relative[str(self.person_number.get())][33])
            self.kozterulet_tipusa3.set(relative[str(self.person_number.get())][34])
            self.kozterulet_tipusa3_egyeb.set(relative[str(self.person_number.get())][35])
            self.address3_V.set(relative[str(self.person_number.get())][36])
            self.megegyezik1.set(relative[str(self.person_number.get())][37])
            self.megegyezik2.set(relative[str(self.person_number.get())][38])
            self.hozzatart_V.set(relative[str(self.person_number.get())][39])
            self.eler_V.set(relative[str(self.person_number.get())][40])
            self.soron_kivul_V.set(relative[str(self.person_number.get())][41])
            self.elogondozás_V.set(relative[str(self.person_number.get())][42])
            self.kdate1_V.set(relative[str(self.person_number.get())][43])
            self.kdate2_V.set(relative[str(self.person_number.get())][44])
            self.kdate3_V.set(relative[str(self.person_number.get())][45])
            self.edate1_V.set(relative[str(self.person_number.get())][46])
            self.edate2_V.set(relative[str(self.person_number.get())][47])
            self.edate3_V.set(relative[str(self.person_number.get())][48])
            self.mothername.set(relative[str(self.person_number.get())][49])
            self.gondnok_nev.set(relative[str(self.person_number.get())][50])
            self.gondnok_cim.set(relative[str(self.person_number.get())][51])
        except:
            pass

    def fill_2(self):
        if  self.forma.get()==self.formak[0]:
            relative=self.ugyfellista
        elif self.forma.get()==self.formak[1]:
            relative = self.ugyfellista_ejjeli
        else:
            relative = self.ugyfellista_nappali
        try:
            self.person_number.set(self.person)
            self.name_var.set(relative[self.person][0])
            self.TAJ1_V.set(relative[self.person][1])
            self.TAJ2_V.set(relative[self.person][2])
            self.TAJ3_V.set(relative[self.person][3])
            self.TAJ4_V.set(relative[self.person][4])
            self.nincsTAJ.set(relative[self.person][5])
            self.bname_var.set(relative[self.person][6])
            self.same_V.set(relative[self.person][7])
            self.nation_var.set(relative[self.person][8])
            self.nationality_var.set(relative[self.person][9])
            self.bplace_V.set(relative[self.person][10])
            self.bdate1_V.set(relative[self.person][11])
            self.bdate2_V.set(relative[self.person][12])
            self.bdate3_V.set(relative[self.person][13])
            self.phone_number1.set(relative[self.person][14])
            self.phone_number2.set(relative[self.person][15])
            self.idegen.set(relative[self.person][16])
            self.cselekves.set(relative[self.person][17])
            self.gondnok.set(relative[self.person][18])
            self.IRSZ_1.set(relative[self.person][19])
            self.city_1.set(relative[self.person][20])
            self.ucca1_V.set(relative[self.person][21])
            self.kozterulet_tipusa.set(relative[self.person][22])
            self.kozterulet_tipusa_egyeb.set(relative[self.person][23])
            self.address_V.set(relative[self.person][24])
            self.IRSZ_2.set(relative[self.person][25])
            self.city_2.set(relative[self.person][26])
            self.ucca2_V.set(relative[self.person][27])
            self.kozterulet_tipusa2.set(relative[self.person][28])
            self.kozterulet_tipusa2_egyeb.set(relative[self.person][29])
            self.address2_V.set(relative[self.person][30])
            self.IRSZ_3.set(relative[self.person][31])
            self.city_3.set(relative[self.person][32])
            self.ucca3_V.set(relative[self.person][33])
            self.kozterulet_tipusa3.set(relative[self.person][34])
            self.kozterulet_tipusa3_egyeb.set(relative[self.person][35])
            self.address3_V.set(relative[self.person][36])
            self.megegyezik1.set(relative[self.person][37])
            self.megegyezik2.set(relative[self.person][38])
            self.hozzatart_V.set(relative[self.person][39])
            self.eler_V.set(relative[self.person][40])
            self.soron_kivul_V.set(relative[self.person][41])
            self.elogondozás_V.set(relative[self.person][42])
            self.kdate1_V.set(relative[self.person][43])
            self.kdate2_V.set(relative[self.person][44])
            self.kdate3_V.set(relative[self.person][45])
            self.edate1_V.set(relative[self.person][46])
            self.edate2_V.set(relative[self.person][47])
            self.edate3_V.set(relative[self.person][48])
            self.mothername.set(relative[self.person][49])
            self.gondnok_nev.set(relative[self.person][50])
            self.gondnok_cim.set(relative[self.person][51])
        except:
            pass

    def zero_record(self):

            self.list_of_people.select_clear(0, END)
            self.active_name=None

            self.name_var.set("")
            self.TAJ1_V.set("")
            self.TAJ2_V.set("")
            self.TAJ3_V.set("")
            self.TAJ4_V.set("")
            self.nincsTAJ.set(0)
            self.bname_var.set("")
            self.same_V.set(1)
            self.nation_var.set(1)
            self.nationality_var.set("")
            self.bplace_V.set("")
            self.bdate1_V.set("")
            self.bdate2_V.set("")
            self.bdate3_V.set("")
            self.phone_number1.set("")
            self.phone_number2.set("")
            self.idegen.set("")
            self.cselekves.set(1)
            self.gondnok.set(0)
            self.IRSZ_1.set("")
            self.city_1.set("")
            self.ucca1_V.set("")
            self.kozterulet_tipusa.set("")
            self.kozterulet_tipusa_egyeb.set("")
            self.address_V.set("")
            self.IRSZ_2.set("")
            self.city_2.set("")
            self.ucca2_V.set("")
            self.kozterulet_tipusa2.set("")
            self.kozterulet_tipusa2_egyeb.set("")
            self.address2_V.set("")
            self.IRSZ_3.set("")
            self.city_3.set("")
            self.ucca3_V.set("")
            self.kozterulet_tipusa3.set("")
            self.kozterulet_tipusa3_egyeb.set("")
            self.address3_V.set("")
            self.megegyezik1.set(1)
            self.megegyezik2.set(1)
            self.hozzatart_V.set("")
            self.eler_V.set("")
            self.soron_kivul_V.set(0)
            self.elogondozás_V.set(0)
            self.edate1_V.set("")
            self.edate2_V.set("")
            self.edate3_V.set("")
            self.kdate1_V.set("")
            self.kdate2_V.set("")
            self.kdate3_V.set("")
            self.list_of_people.select_clear(0, END)
            self.mothername.set("")
            self.gondnok_nev.set("")
            self.gondnok_cim.set("")
            self.list_of_people.select_clear(0, END)

            self.active_name=None
            self.list_of_people.select_clear(0, END)
            self.list_of_people.select_clear(0, END)
            self.list_of_people.delete(0, END)

            if str(self.forma.get()) == self.formak[0]:
                self.person_number_new = (len(self.ugyfellista) + 1)

                try:
                    for person in self.ugyfellista:
                        self.list_of_people.insert(END,
                                                   str(person) + ".: " + str(self.ugyfellista[person][0]) + " (" + str(
                                                       self.ugyfellista[person][11]) + ")")
                except:
                    pass
            elif str(self.forma.get()) == self.formak[1]:
                self.person_number_new = (len(self.ugyfellista_ejjeli) + 1)
                try:
                    for person in self.ugyfellista_ejjeli:
                        self.list_of_people.insert(END, str(person) + ".: " + str(
                            self.ugyfellista_ejjeli[person][0]) + " (" + str(self.ugyfellista_ejjeli[person][11]) + ")")
                except:
                    pass
            else:
                self.person_number_new = (len(self.ugyfellista_nappali) + 1)
                try:
                    for person in self.ugyfellista_nappali:
                        self.list_of_people.insert(END, str(person) + ".: " + str(
                            self.ugyfellista_nappali[person][0]) + " (" + str(
                            self.ugyfellista_nappali[person][11]) + ")")
                except:
                    pass
            self.person_number.set(self.person_number_new)

if __name__ == "__main__":
    Ablak = Tk()
    Ablak.title("Szt. 20.§ szerinti 'Nyilvántartás' Hajléktalanok Átmenti Szállása részére")
    Ablak.geometry("825x685")
    Ablak.resizable(False,False)
    Felul = Felulet(Ablak)
    Ablak.mainloop()

